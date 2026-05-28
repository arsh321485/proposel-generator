# # """
# # main.py  —  SecureITLab Proposal Generator Backend
# # Run with:  uvicorn main:app --reload --port 8000
# # """

# # from fastapi import FastAPI, HTTPException
# # from fastapi.middleware.cors import CORSMiddleware
# # from fastapi.responses import StreamingResponse, JSONResponse
# # from fastapi.staticfiles import StaticFiles
# # from pydantic import BaseModel
# # from typing import List, Optional
# # import io, os

# # from docx import Document
# # from docx.shared import Pt, RGBColor, Cm, Inches
# # from docx.enum.text import WD_ALIGN_PARAGRAPH
# # from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
# # from docx.oxml.ns import qn
# # from docx.oxml import OxmlElement
# # from lxml import etree

# # from proposal_data import PROPOSALS

# # # ── Colour helpers ──────────────────────────────────────────────────────────
# # NAVY   = RGBColor(0x1A, 0x3C, 0x6E)
# # BLUE   = RGBColor(0x2E, 0x75, 0xB6)
# # LBLUE  = RGBColor(0xD5, 0xE8, 0xF0)
# # GREY   = RGBColor(0x59, 0x59, 0x59)
# # WHITE  = RGBColor(0xFF, 0xFF, 0xFF)

# # NAVY_HEX  = "1A3C6E"
# # BLUE_HEX  = "2E75B6"
# # LBLUE_HEX = "D5E8F0"
# # GREY_HEX  = "595959"
# # WHITE_HEX = "FFFFFF"
# # CC_HEX    = "CCCCCC"

# # # ── Pydantic models ─────────────────────────────────────────────────────────
# # class CostRow(BaseModel):
# #     item: str = ""
# #     desc: str = ""
# #     unit_price: float = 0
# #     qty: int = 1
# #     total: float = 0

# # class TimelineRow(BaseModel):
# #     phase: str = ""
# #     activities: str = ""
# #     duration: str = ""
# #     deliverables: str = ""

# # class ProposalRequest(BaseModel):
# #     # Service
# #     service_key: str                     # key from PROPOSALS dict
# #     service_name: str
# #     service_desc: str
# #     service_deliverables: List[str] = []

# #     # Organisation
# #     org_name: str = ""
# #     org_industry: str = ""
# #     org_address: str = ""
# #     org_contact: str = ""
# #     org_email: str = ""
# #     org_phone: str = ""
# #     org_bg: str = ""
# #     org_obj: str = ""

# #     # Proposal meta
# #     prop_date: str = ""
# #     prop_ref: str = ""
# #     start_date: str = ""
# #     end_date: str = ""
# #     currency: str = "USD"
# #     tax_rate: float = 0

# #     # Scope lists
# #     inscope: List[str] = []
# #     outscope: List[str] = []
# #     assumptions: List[str] = []
# #     responsibilities: List[str] = []
# #     payments: List[str] = []

# #     # Tables (user-customised; if empty, use proposal_data defaults)
# #     timeline_rows: List[TimelineRow] = []
# #     cost_rows: List[CostRow] = []

# # # ── FastAPI app ─────────────────────────────────────────────────────────────
# # app = FastAPI(title="SecureITLab Proposal Generator", version="1.0.0")

# # app.add_middleware(
# #     CORSMiddleware,
# #     allow_origins=["*"],          # tighten for production
# #     allow_methods=["*"],
# #     allow_headers=["*"],
# # )

# # # Serve the HTML front-end at /
# # if os.path.exists("SecureITLab_Proposal_Generator.html"):
# #     from fastapi.responses import FileResponse
# #     @app.get("/")
# #     def serve_frontend():
# #         return FileResponse("SecureITLab_Proposal_Generator.html")

# # # ── Endpoint: list available services ───────────────────────────────────────
# # @app.get("/services")
# # def list_services():
# #     return {k: {"phases": len(v["timeline"])} for k, v in PROPOSALS.items()}

# # # ── Endpoint: get methodology + timeline for a service ───────────────────────
# # @app.get("/services/{key}")
# # def get_service(key: str):
# #     if key not in PROPOSALS:
# #         raise HTTPException(404, f"Service key '{key}' not found")
# #     data = PROPOSALS[key]
# #     return {
# #         "methodology": data["methodology"],
# #         "timeline": [
# #             {"phase": t[0], "activities": t[1], "duration": t[2], "deliverables": t[3]}
# #             for t in data["timeline"]
# #         ]
# #     }

# # # ── Endpoint: generate proposal docx ────────────────────────────────────────
# # @app.post("/generate")
# # def generate_proposal(req: ProposalRequest):
# #     if req.service_key not in PROPOSALS:
# #         raise HTTPException(404, f"Service key '{req.service_key}' not found")

# #     proposal_data = PROPOSALS[req.service_key]

# #     # Use timeline from proposal_data if user didn't customise
# #     if not req.timeline_rows:
# #         timeline = [
# #             TimelineRow(phase=t[0], activities=t[1], duration=t[2], deliverables=t[3])
# #             for t in proposal_data["timeline"]
# #         ]
# #     else:
# #         timeline = req.timeline_rows

# #     buf = build_docx(req, proposal_data, timeline)

# #     safe_name = (req.org_name or "Proposal").replace(" ", "_")
# #     safe_svc  = req.service_name.replace(" ", "_").replace("/", "-")
# #     filename  = f"Proposal_{safe_name}_{safe_svc}.docx"

# #     return StreamingResponse(
# #         buf,
# #         media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
# #         headers={"Content-Disposition": f'attachment; filename="{filename}"'},
# #     )

# # # ── Document builder ─────────────────────────────────────────────────────────
# # def build_docx(req: ProposalRequest, proposal_data: dict, timeline: list) -> io.BytesIO:
# #     doc = Document()

# #     # ── Page setup (A4, 2.5 cm margins) ─────────────────────────────────────
# #     for section in doc.sections:
# #         section.page_width  = Cm(21)
# #         section.page_height = Cm(29.7)
# #         section.top_margin = section.bottom_margin = Cm(2.5)
# #         section.left_margin = section.right_margin = Cm(2.5)

# #     # ── Document-wide default font ───────────────────────────────────────────
# #     style = doc.styles["Normal"]
# #     style.font.name = "Arial"
# #     style.font.size = Pt(10)

# #     # ── Header & footer ──────────────────────────────────────────────────────
# #     hdr = doc.sections[0].header
# #     hdr_para = hdr.paragraphs[0]
# #     hdr_para.clear()
# #     _add_run(hdr_para, f"SecureITLab — Confidential     |     {req.service_name}",
# #              size=8, colour=GREY)
# #     _bottom_border(hdr_para, BLUE_HEX, 6)

# #     ftr = doc.sections[0].footer
# #     ftr_para = ftr.paragraphs[0]
# #     ftr_para.clear()
# #     _add_run(ftr_para, f"www.secureitlab.com     |     {req.org_name}     |     {req.prop_date}",
# #              size=8, colour=GREY)
# #     _top_border(ftr_para, CC_HEX, 4)

# #     # ══════════════════════════════════════════════════════════════════
# #     # COVER PAGE
# #     # ══════════════════════════════════════════════════════════════════
# #     _spacer(doc, 6)
# #     p = doc.add_paragraph()
# #     p.alignment = WD_ALIGN_PARAGRAPH.CENTER
# #     _add_run(p, "SecureITLab", size=28, bold=True, colour=NAVY)

# #     p = doc.add_paragraph()
# #     p.alignment = WD_ALIGN_PARAGRAPH.CENTER
# #     _add_run(p, "Cybersecurity  |  Data Privacy  |  Data Governance", size=11, colour=BLUE)
# #     _bottom_border(p, BLUE_HEX, 6)

# #     _spacer(doc, 3)
# #     p = doc.add_paragraph()
# #     p.alignment = WD_ALIGN_PARAGRAPH.CENTER
# #     _add_run(p, "SERVICE PROPOSAL", size=14, bold=True, colour=GREY)

# #     p = doc.add_paragraph()
# #     p.alignment = WD_ALIGN_PARAGRAPH.CENTER
# #     _add_run(p, req.service_name, size=18, bold=True, colour=NAVY)
# #     _bottom_border(p, BLUE_HEX, 6)

# #     _spacer(doc, 2)
# #     p = doc.add_paragraph()
# #     p.alignment = WD_ALIGN_PARAGRAPH.CENTER
# #     _add_run(p, "Prepared for:", size=11, colour=GREY)

# #     p = doc.add_paragraph()
# #     p.alignment = WD_ALIGN_PARAGRAPH.CENTER
# #     _add_run(p, req.org_name, size=16, bold=True, colour=NAVY)

# #     _spacer(doc, 1)
# #     p = doc.add_paragraph()
# #     p.alignment = WD_ALIGN_PARAGRAPH.CENTER
# #     _add_run(p, f"Proposal Date: {req.prop_date}", size=10, colour=GREY)

# #     p = doc.add_paragraph()
# #     p.alignment = WD_ALIGN_PARAGRAPH.CENTER
# #     _add_run(p, f"Reference: {req.prop_ref}", size=10, colour=GREY)

# #     doc.add_page_break()

# #     # ══════════════════════════════════════════════════════════════════
# #     # S1  Organisation Overview
# #     # ══════════════════════════════════════════════════════════════════
# #     _heading(doc, "1.  Organisation Overview")
# #     _rule(doc)

# #     fields = [
# #         ("Organisation Name", req.org_name),
# #         ("Industry / Sector",  req.org_industry),
# #         ("Registered Address", req.org_address),
# #         ("Primary Contact",    req.org_contact),
# #         ("Email",              req.org_email),
# #         ("Phone",              req.org_phone),
# #         ("Organisation Background", req.org_bg),
# #         ("Engagement Objective",    req.org_obj),
# #     ]
# #     for label, val in fields:
# #         p = doc.add_paragraph()
# #         _add_run(p, f"{label}:  ", bold=True, colour=NAVY)
# #         _add_run(p, val or "—", colour=GREY)

# #     doc.add_page_break()

# #     # ══════════════════════════════════════════════════════════════════
# #     # S2  About SecureITLab  (boilerplate)
# #     # ══════════════════════════════════════════════════════════════════
# #     _heading(doc, "2.  About SecureITLab")
# #     _rule(doc)

# #     _body(doc, ("SecureITLab is a remote-first cybersecurity, data privacy, and data governance "
# #                 "consulting firm with over 20 years of industry experience. Headquartered in the "
# #                 "Kingdom of Bahrain, with offices in the Philippines and India, SecureITLab serves "
# #                 "clients across the Middle East, Asia-Pacific, and beyond."))
# #     _body(doc, ("Our team of certified professionals (CISSP, CISM, CISA, CEH, ISO 27001 LA/LI, CDPSE) "
# #                 "delivers tailored, risk-driven solutions that empower organisations to navigate the "
# #                 "complex landscape of cybersecurity and compliance with confidence."))
# #     _body(doc, "SecureITLab is certified to ISO 27001:2022, ISO 27701:2019, and ISO 22301:2019.")

# #     _subheading(doc, "Core Capabilities")
# #     caps = [
# #         "Consulting — Cybersecurity and Data Privacy",
# #         "Assurance — Penetration Testing, Vulnerability Assessment and Red Teaming",
# #         "Audit and Compliance — IT Audit, Cloud Audit and ISO Compliance",
# #         "Services Augmentation — vCISO, vCDO and Expert Resourcing",
# #         "Managed Services — SOC, NOC, Vulnerability Management and Take-Down",
# #     ]
# #     for c in caps:
# #         _bullet(doc, c)

# #     doc.add_page_break()

# #     # ══════════════════════════════════════════════════════════════════
# #     # S3  Service Description
# #     # ══════════════════════════════════════════════════════════════════
# #     _heading(doc, "3.  Service Description")
# #     _rule(doc)

# #     p = doc.add_paragraph()
# #     _add_run(p, req.service_name, size=12, bold=True, colour=BLUE)
# #     _body(doc, req.service_desc)

# #     _subheading(doc, "Key Deliverables")
# #     for d in req.service_deliverables:
# #         _bullet(doc, d)

# #     _subheading(doc, "Methodology")
# #     _body(doc, ("SecureITLab employs industry-recognised methodologies, standards, and frameworks "
# #                 "throughout this engagement. All activities are conducted in alignment with the agreed "
# #                 "scope and subject to the client's change-control and information security policies."))
# #     for phase_text in proposal_data["methodology"]:
# #         _body(doc, phase_text)

# #     doc.add_page_break()

# #     # ══════════════════════════════════════════════════════════════════
# #     # S4  Scope of Work
# #     # ══════════════════════════════════════════════════════════════════
# #     _heading(doc, "4.  Scope of Work")
# #     _rule(doc)
# #     _body(doc, "Define the precise boundaries of the engagement. Items not listed are considered out of scope.")

# #     _subheading(doc, "In Scope")
# #     _bullet_list(doc, req.inscope, fallback="To be defined.")

# #     _subheading(doc, "Out of Scope")
# #     _bullet_list(doc, req.outscope, fallback="To be defined.")

# #     _subheading(doc, "Assumptions and Dependencies")
# #     _bullet_list(doc, req.assumptions, fallback="To be confirmed.")

# #     _subheading(doc, "Client Responsibilities")
# #     _bullet_list(doc, req.responsibilities, fallback="To be confirmed.")

# #     doc.add_page_break()

# #     # ══════════════════════════════════════════════════════════════════
# #     # S5  Project Timeline
# #     # ══════════════════════════════════════════════════════════════════
# #     _heading(doc, "5.  Project Timeline")
# #     _rule(doc)
# #     _body(doc, "The indicative timeline below is subject to confirmation upon contract award.")

# #     p = doc.add_paragraph()
# #     _add_run(p, "Proposed Start: ", bold=True, colour=NAVY)
# #     _add_run(p, req.start_date or "TBD", colour=GREY)
# #     _add_run(p, "     Proposed End: ", bold=True, colour=NAVY)
# #     _add_run(p, req.end_date or "TBD", colour=GREY)

# #     # Timeline table
# #     col_w = [Cm(3.2), Cm(6.2), Cm(2.8), Cm(5.5)]
# #     tbl = doc.add_table(rows=1, cols=4)
# #     tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
# #     _set_col_widths(tbl, col_w)

# #     hdr_cells = tbl.rows[0].cells
# #     for i, h in enumerate(["Phase", "Activities", "Duration", "Deliverables"]):
# #         _header_cell(hdr_cells[i], h, col_w[i])

# #     for idx, row in enumerate(timeline):
# #         cells = tbl.add_row().cells
# #         fill = LBLUE_HEX if idx % 2 == 1 else WHITE_HEX
# #         _data_cell(cells[0], row.phase,        col_w[0], fill)
# #         _data_cell(cells[1], row.activities,   col_w[1], fill)
# #         _data_cell(cells[2], row.duration,     col_w[2], fill)
# #         _data_cell(cells[3], row.deliverables, col_w[3], fill)

# #     doc.add_page_break()

# #     # ══════════════════════════════════════════════════════════════════
# #     # S6  Investment Summary
# #     # ══════════════════════════════════════════════════════════════════
# #     _heading(doc, "6.  Investment Summary")
# #     _rule(doc)
# #     _body(doc, "All fees are exclusive of applicable taxes. Payment terms are as per the Master Service Agreement or as agreed in writing.")

# #     p = doc.add_paragraph()
# #     _add_run(p, f"Currency: {req.currency}", bold=True, colour=NAVY)

# #     cost_rows = req.cost_rows
# #     subtotal  = sum(r.total for r in cost_rows)
# #     tax_amt   = round(subtotal * req.tax_rate / 100, 2)
# #     grand     = round(subtotal + tax_amt, 2)

# #     ccol = [Cm(4.2), Cm(5.4), Cm(2.8), Cm(1.6), Cm(3.7)]
# #     ctbl = doc.add_table(rows=1, cols=5)
# #     ctbl.alignment = WD_TABLE_ALIGNMENT.LEFT
# #     _set_col_widths(ctbl, ccol)

# #     ch = ctbl.rows[0].cells
# #     for i, h in enumerate(["Line Item", "Description", "Unit Price", "Qty", f"Total ({req.currency})"]):
# #         _header_cell(ch[i], h, ccol[i])

# #     for idx, r in enumerate(cost_rows):
# #         cells = ctbl.add_row().cells
# #         fill  = LBLUE_HEX if idx % 2 == 1 else WHITE_HEX
# #         _data_cell(cells[0], r.item,             ccol[0], fill)
# #         _data_cell(cells[1], r.desc,             ccol[1], fill)
# #         _data_cell(cells[2], _fmt(r.unit_price), ccol[2], fill)
# #         _data_cell(cells[3], str(r.qty),         ccol[3], fill)
# #         _data_cell(cells[4], _fmt(r.total),      ccol[4], fill)

# #     # Subtotal / tax / grand total rows
# #     _summary_row(ctbl, ccol, f"Subtotal",           _fmt(subtotal), LBLUE_HEX, NAVY_HEX)
# #     _summary_row(ctbl, ccol, f"Tax ({req.tax_rate}%)", _fmt(tax_amt), LBLUE_HEX, NAVY_HEX)
# #     _summary_row(ctbl, ccol, "TOTAL INVESTMENT",
# #                  f"{req.currency} {_fmt(grand)}", NAVY_HEX, WHITE_HEX, bold_label=True)

# #     _subheading(doc, "Payment Schedule")
# #     _bullet_list(doc, req.payments, fallback="As agreed.")
# #     _body(doc, "Validity: This proposal is valid for 30 days from the date of issue.")

# #     doc.add_page_break()

# #     # ══════════════════════════════════════════════════════════════════
# #     # S7  Terms and Conditions  (boilerplate)
# #     # ══════════════════════════════════════════════════════════════════
# #     _heading(doc, "7.  Terms and Conditions")
# #     _rule(doc)
# #     _body(doc, ("This proposal is subject to SecureITLab's standard Master Service Agreement (MSA). "
# #                 "Key terms are summarised below; the MSA governs in the event of any conflict."))

# #     tc_items = [
# #         ("Confidentiality",
# #          "Both parties agree to maintain strict confidentiality of all information exchanged. "
# #          "SecureITLab will not disclose client data to third parties without prior written consent."),
# #         ("Intellectual Property",
# #          "All deliverables become the property of the client upon receipt of full payment. "
# #          "SecureITLab retains rights to its pre-existing methodologies, tools, and frameworks."),
# #         ("Limitation of Liability",
# #          "SecureITLab's aggregate liability shall not exceed the total fees paid under this engagement. "
# #          "SecureITLab is not liable for indirect, consequential, or incidental damages."),
# #         ("Governing Law",
# #          "This proposal and any resulting engagement shall be governed by the laws of the Kingdom of "
# #          "Bahrain, unless otherwise agreed in writing."),
# #         ("Amendments",
# #          "Any change in scope, timeline, or cost requires a written Change Request signed by both "
# #          "parties prior to implementation."),
# #     ]
# #     for title, body in tc_items:
# #         p = doc.add_paragraph()
# #         _add_run(p, f"{title}:  ", bold=True, colour=NAVY)
# #         _add_run(p, body, colour=GREY)

# #     doc.add_page_break()

# #     # ══════════════════════════════════════════════════════════════════
# #     # S8  Acceptance and Signatures
# #     # ══════════════════════════════════════════════════════════════════
# #     _heading(doc, "8.  Acceptance and Signatures")
# #     _rule(doc)
# #     _body(doc, ("By signing below, both parties confirm their agreement to the terms set out in this "
# #                 "proposal and authorise SecureITLab to commence the engagement as described."))

# #     _spacer(doc, 2)

# #     sig_col = [Cm(7.5), Cm(2.0), Cm(7.5)]
# #     stbl = doc.add_table(rows=4, cols=3)
# #     _set_col_widths(stbl, sig_col)

# #     header_row = stbl.rows[0]
# #     _sig_header_cell(header_row.cells[0], f"Authorised Signatory — {req.org_name}", sig_col[0])
# #     _sig_header_cell(header_row.cells[2], "Authorised Signatory — SecureITLab",     sig_col[2])

# #     for row in stbl.rows[1:]:
# #         for cell in row.cells:
# #             cell.text = ""

# #     _spacer(doc, 2)
# #     _rule(doc)

# #     p = doc.add_paragraph()
# #     p.alignment = WD_ALIGN_PARAGRAPH.CENTER
# #     _add_run(p, "SecureITLab  |  info@secureitlab.com  |  www.secureitlab.com", size=9, colour=GREY)

# #     p = doc.add_paragraph()
# #     p.alignment = WD_ALIGN_PARAGRAPH.CENTER
# #     _add_run(p, "Bahrain  |  Philippines  |  India", size=9, colour=GREY)

# #     buf = io.BytesIO()
# #     doc.save(buf)
# #     buf.seek(0)
# #     return buf


# # # ── Low-level helpers ────────────────────────────────────────────────────────

# # def _add_run(para, text, size=10, bold=False, italic=False, colour=None):
# #     run = para.add_run(text)
# #     run.font.name = "Arial"
# #     run.font.size = Pt(size)
# #     run.font.bold = bold
# #     run.font.italic = italic
# #     if colour:
# #         run.font.color.rgb = colour
# #     return run

# # def _heading(doc, text):
# #     p = doc.add_paragraph()
# #     p.paragraph_format.space_before = Pt(14)
# #     p.paragraph_format.space_after  = Pt(6)
# #     _add_run(p, text, size=14, bold=True, colour=NAVY)

# # def _subheading(doc, text):
# #     p = doc.add_paragraph()
# #     p.paragraph_format.space_before = Pt(8)
# #     p.paragraph_format.space_after  = Pt(4)
# #     _add_run(p, text, size=11, bold=True, colour=NAVY)

# # def _body(doc, text):
# #     p = doc.add_paragraph()
# #     p.paragraph_format.space_before = Pt(2)
# #     p.paragraph_format.space_after  = Pt(4)
# #     _add_run(p, text, colour=GREY)

# # def _bullet(doc, text):
# #     p = doc.add_paragraph(style="List Bullet")
# #     _add_run(p, text, colour=GREY)

# # def _bullet_list(doc, items, fallback=""):
# #     clean = [i for i in items if i.strip()]
# #     if clean:
# #         for item in clean:
# #             _bullet(doc, item)
# #     elif fallback:
# #         _body(doc, fallback)

# # def _spacer(doc, lines=1):
# #     for _ in range(lines):
# #         doc.add_paragraph()

# # def _rule(doc):
# #     p = doc.add_paragraph()
# #     _bottom_border(p, BLUE_HEX, 6)

# # def _fmt(val):
# #     return f"{val:,.2f}"

# # def _bottom_border(para, colour_hex, size):
# #     pPr = para._p.get_or_add_pPr()
# #     pBdr = OxmlElement("w:pBdr")
# #     bottom = OxmlElement("w:bottom")
# #     bottom.set(qn("w:val"), "single")
# #     bottom.set(qn("w:sz"), str(size))
# #     bottom.set(qn("w:space"), "1")
# #     bottom.set(qn("w:color"), colour_hex)
# #     pBdr.append(bottom)
# #     pPr.append(pBdr)

# # def _top_border(para, colour_hex, size):
# #     pPr = para._p.get_or_add_pPr()
# #     pBdr = OxmlElement("w:pBdr")
# #     top = OxmlElement("w:top")
# #     top.set(qn("w:val"), "single")
# #     top.set(qn("w:sz"), str(size))
# #     top.set(qn("w:space"), "1")
# #     top.set(qn("w:color"), colour_hex)
# #     pBdr.append(top)
# #     pPr.append(pBdr)

# # def _set_col_widths(table, widths):
# #     tbl_el = table._tbl
# #     tblPr  = tbl_el.find(qn("w:tblPr")) or OxmlElement("w:tblPr")
# #     tblW   = OxmlElement("w:tblW")
# #     total  = sum(int(w.cm * 914400 / 360) for w in widths)   # EMU → twips approx
# #     tblW.set(qn("w:w"),    str(total))
# #     tblW.set(qn("w:type"), "dxa")
# #     tblPr.append(tblW)

# # def _twips(cm_val):
# #     return int(cm_val.cm * 567)   # 1 cm ≈ 567 twips

# # def _shade_cell(cell, fill_hex):
# #     tcPr = cell._tc.get_or_add_tcPr()
# #     shd  = OxmlElement("w:shd")
# #     shd.set(qn("w:val"),   "clear")
# #     shd.set(qn("w:color"), "auto")
# #     shd.set(qn("w:fill"),  fill_hex)
# #     tcPr.append(shd)

# # def _cell_borders(cell, colour_hex, size):
# #     tcPr = cell._tc.get_or_add_tcPr()
# #     tcBorders = OxmlElement("w:tcBorders")
# #     for side in ("top", "left", "bottom", "right"):
# #         el = OxmlElement(f"w:{side}")
# #         el.set(qn("w:val"),   "single")
# #         el.set(qn("w:sz"),    str(size))
# #         el.set(qn("w:color"), colour_hex)
# #         tcBorders.append(el)
# #     tcPr.append(tcBorders)

# # def _cell_width(cell, cm_val):
# #     tcPr = cell._tc.get_or_add_tcPr()
# #     tcW  = OxmlElement("w:tcW")
# #     tcW.set(qn("w:w"),    str(_twips(cm_val)))
# #     tcW.set(qn("w:type"), "dxa")
# #     tcPr.append(tcW)

# # def _header_cell(cell, text, width):
# #     cell.text = ""
# #     _shade_cell(cell, NAVY_HEX)
# #     _cell_borders(cell, BLUE_HEX, 6)
# #     _cell_width(cell, width)
# #     p = cell.paragraphs[0]
# #     p.alignment = WD_ALIGN_PARAGRAPH.CENTER
# #     _add_run(p, text, size=9, bold=True, colour=WHITE)

# # def _data_cell(cell, text, width, fill=WHITE_HEX):
# #     cell.text = ""
# #     _shade_cell(cell, fill)
# #     _cell_borders(cell, CC_HEX, 4)
# #     _cell_width(cell, width)
# #     p = cell.paragraphs[0]
# #     _add_run(p, text or "", size=9, colour=GREY)

# # def _summary_row(table, col_widths, label_text, value_text, label_fill, label_colour_hex, bold_label=False):
# #     row   = table.add_row()
# #     cells = row.cells
# #     # Merge first 4 cells for label
# #     merged = cells[0].merge(cells[1]).merge(cells[2]).merge(cells[3])
# #     _shade_cell(merged, label_fill)
# #     _cell_borders(merged, BLUE_HEX, 6)
# #     p = merged.paragraphs[0]
# #     p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
# #     lc = WHITE if label_fill == NAVY_HEX else NAVY
# #     _add_run(p, label_text, size=10, bold=bold_label, colour=lc)

# #     vc = cells[4]
# #     _shade_cell(vc, label_fill)
# #     _cell_borders(vc, BLUE_HEX, 6)
# #     _cell_width(vc, col_widths[4])
# #     vp = vc.paragraphs[0]
# #     vc_text_colour = WHITE if label_fill == NAVY_HEX else GREY
# #     _add_run(vp, value_text, size=10, bold=bold_label, colour=vc_text_colour)

# # def _sig_header_cell(cell, text, width):
# #     _cell_width(cell, width)
# #     tcPr = cell._tc.get_or_add_tcPr()
# #     tcBorders = OxmlElement("w:tcBorders")
# #     bottom = OxmlElement("w:bottom")
# #     bottom.set(qn("w:val"),   "single")
# #     bottom.set(qn("w:sz"),    "6")
# #     bottom.set(qn("w:color"), NAVY_HEX)
# #     tcBorders.append(bottom)
# #     tcPr.append(tcBorders)
# #     p = cell.paragraphs[0]
# #     p.paragraph_format.space_before = Pt(40)
# #     _add_run(p, text, size=9, bold=True, colour=NAVY)
























# """
# main.py  —  SecureITLab Proposal Generator Backend
# Run with:  uvicorn main:app --reload --port 8000

# NEW in this version:
#   • Every /generate call is logged to downloads.csv
#   • Admin dashboard at /dashboard  (login required)
#   • Admin credentials set via ADMIN_USER / ADMIN_PASS env vars
#     (defaults: admin / secureitlab2024)
# """

# from fastapi import FastAPI, HTTPException, Request, Depends
# from fastapi.middleware.cors import CORSMiddleware
# from fastapi.responses import StreamingResponse, JSONResponse, HTMLResponse, RedirectResponse
# from fastapi.security import HTTPBasic, HTTPBasicCredentials
# from pydantic import BaseModel
# from typing import List, Optional
# import io, os, csv, secrets
# from datetime import datetime
# from pathlib import Path

# from docx import Document
# from docx.shared import Pt, RGBColor, Cm, Inches
# from docx.enum.text import WD_ALIGN_PARAGRAPH
# from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
# from docx.oxml.ns import qn
# from docx.oxml import OxmlElement
# from lxml import etree

# from proposal_data import PROPOSALS

# # ── Admin credentials (set env vars to override) ─────────────────────────────
# ADMIN_USER = os.getenv("ADMIN_USER", "admin")
# ADMIN_PASS = os.getenv("ADMIN_PASS", "secureitlab2024")

# # ── CSV log file ──────────────────────────────────────────────────────────────
# LOG_FILE = Path("downloads.csv")
# CSV_HEADERS = ["Timestamp", "Client Name", "Organisation", "Service", "Reference", "IP Address", "Filename"]

# def _ensure_log():
#     """Create the CSV with headers if it doesn't exist."""
#     if not LOG_FILE.exists():
#         with open(LOG_FILE, "w", newline="", encoding="utf-8") as f:
#             writer = csv.writer(f)
#             writer.writerow(CSV_HEADERS)

# def _log_download(req_data, ip: str, filename: str):
#     """Append one row to downloads.csv."""
#     _ensure_log()
#     with open(LOG_FILE, "a", newline="", encoding="utf-8") as f:
#         writer = csv.writer(f)
#         writer.writerow([
#             datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
#             req_data.org_contact or "—",
#             req_data.org_name    or "—",
#             req_data.service_name,
#             req_data.prop_ref    or "—",
#             ip,
#             filename,
#         ])

# def _read_log() -> list:
#     """Return all rows from the CSV as a list of dicts."""
#     _ensure_log()
#     with open(LOG_FILE, "r", encoding="utf-8") as f:
#         return list(csv.DictReader(f))

# # ── Colour helpers ────────────────────────────────────────────────────────────
# NAVY   = RGBColor(0x1A, 0x3C, 0x6E)
# BLUE   = RGBColor(0x2E, 0x75, 0xB6)
# LBLUE  = RGBColor(0xD5, 0xE8, 0xF0)
# GREY   = RGBColor(0x59, 0x59, 0x59)
# WHITE  = RGBColor(0xFF, 0xFF, 0xFF)

# NAVY_HEX  = "1A3C6E"
# BLUE_HEX  = "2E75B6"
# LBLUE_HEX = "D5E8F0"
# GREY_HEX  = "595959"
# WHITE_HEX = "FFFFFF"
# CC_HEX    = "CCCCCC"

# # ── Pydantic models ───────────────────────────────────────────────────────────
# class CostRow(BaseModel):
#     item: str = ""
#     desc: str = ""
#     unit_price: float = 0
#     qty: int = 1
#     total: float = 0

# class TimelineRow(BaseModel):
#     phase: str = ""
#     activities: str = ""
#     duration: str = ""
#     deliverables: str = ""

# class ProposalRequest(BaseModel):
#     # Service
#     service_key: str
#     service_name: str
#     service_desc: str
#     service_deliverables: List[str] = []

#     # Organisation
#     org_name: str = ""
#     org_industry: str = ""
#     org_address: str = ""
#     org_contact: str = ""
#     org_email: str = ""
#     org_phone: str = ""
#     org_bg: str = ""
#     org_obj: str = ""

#     # Proposal meta
#     prop_date: str = ""
#     prop_ref: str = ""
#     start_date: str = ""
#     end_date: str = ""
#     currency: str = "USD"
#     tax_rate: float = 0

#     # Scope lists
#     inscope: List[str] = []
#     outscope: List[str] = []
#     assumptions: List[str] = []
#     responsibilities: List[str] = []
#     payments: List[str] = []

#     # Tables
#     timeline_rows: List[TimelineRow] = []
#     cost_rows: List[CostRow] = []

# # ── FastAPI app ───────────────────────────────────────────────────────────────
# app = FastAPI(title="SecureITLab Proposal Generator", version="2.0.0")
# security = HTTPBasic()

# app.add_middleware(
#     CORSMiddleware,
#     allow_origins=["*"],
#     allow_methods=["*"],
#     allow_headers=["*"],
# )

# # Serve the HTML front-end at /
# if os.path.exists("SecureITLab_Proposal_Generator.html"):
#     from fastapi.responses import FileResponse
#     @app.get("/")
#     def serve_frontend():
#         return FileResponse("SecureITLab_Proposal_Generator.html")

# # ── Admin auth helper ─────────────────────────────────────────────────────────
# def _verify_admin(credentials: HTTPBasicCredentials = Depends(security)):
#     ok_user = secrets.compare_digest(credentials.username.encode(), ADMIN_USER.encode())
#     ok_pass = secrets.compare_digest(credentials.password.encode(), ADMIN_PASS.encode())
#     if not (ok_user and ok_pass):
#         raise HTTPException(
#             status_code=401,
#             detail="Invalid credentials",
#             headers={"WWW-Authenticate": "Basic"},
#         )
#     return credentials.username

# # ── Endpoint: list services ───────────────────────────────────────────────────
# @app.get("/services")
# def list_services():
#     return {k: {"phases": len(v["timeline"])} for k, v in PROPOSALS.items()}

# # ── Endpoint: get service detail ──────────────────────────────────────────────
# @app.get("/services/{key}")
# def get_service(key: str):
#     if key not in PROPOSALS:
#         raise HTTPException(404, f"Service key '{key}' not found")
#     data = PROPOSALS[key]
#     return {
#         "methodology": data["methodology"],
#         "timeline": [
#             {"phase": t[0], "activities": t[1], "duration": t[2], "deliverables": t[3]}
#             for t in data["timeline"]
#         ]
#     }

# # ── Endpoint: generate proposal docx ─────────────────────────────────────────
# @app.post("/generate")
# def generate_proposal(req: ProposalRequest, request: Request):
#     if req.service_key not in PROPOSALS:
#         raise HTTPException(404, f"Service key '{req.service_key}' not found")

#     proposal_data = PROPOSALS[req.service_key]

#     if not req.timeline_rows:
#         timeline = [
#             TimelineRow(phase=t[0], activities=t[1], duration=t[2], deliverables=t[3])
#             for t in proposal_data["timeline"]
#         ]
#     else:
#         timeline = req.timeline_rows

#     buf = build_docx(req, proposal_data, timeline)

#     safe_name = (req.org_name or "Proposal").replace(" ", "_")
#     safe_svc  = req.service_name.replace(" ", "_").replace("/", "-")
#     filename  = f"Proposal_{safe_name}_{safe_svc}.docx"

#     # ── Log the download ──────────────────────────────────────────────────────
#     client_ip = request.headers.get("X-Forwarded-For", request.client.host)
#     _log_download(req, client_ip, filename)

#     return StreamingResponse(
#         buf,
#         media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
#         headers={"Content-Disposition": f'attachment; filename="{filename}"'},
#     )

# # ── Endpoint: admin dashboard ─────────────────────────────────────────────────
# @app.get("/dashboard", response_class=HTMLResponse)
# def dashboard(admin: str = Depends(_verify_admin)):
#     rows = _read_log()
#     total      = len(rows)
#     today_str  = datetime.now().strftime("%Y-%m-%d")
#     today_count = sum(1 for r in rows if r["Timestamp"].startswith(today_str))

#     # Count by service
#     service_counts: dict = {}
#     for r in rows:
#         svc = r.get("Service", "Unknown")
#         service_counts[svc] = service_counts.get(svc, 0) + 1
#     top_services = sorted(service_counts.items(), key=lambda x: x[1], reverse=True)[:5]

#     # Build table rows HTML (newest first)
#     table_rows_html = ""
#     for i, r in enumerate(reversed(rows)):
#         bg = "#f8faff" if i % 2 == 0 else "#ffffff"
#         table_rows_html += f"""
#         <tr style="background:{bg};">
#           <td>{r.get('Timestamp','—')}</td>
#           <td>{r.get('Client Name','—')}</td>
#           <td>{r.get('Organisation','—')}</td>
#           <td><span class="badge">{r.get('Service','—')}</span></td>
#           <td>{r.get('Reference','—')}</td>
#           <td>{r.get('IP Address','—')}</td>
#           <td style="font-size:11px;color:#888;">{r.get('Filename','—')}</td>
#         </tr>"""

#     # Build top services HTML
#     top_html = ""
#     for svc, count in top_services:
#         pct = round(count / total * 100) if total else 0
#         top_html += f"""
#         <div style="margin-bottom:10px;">
#           <div style="display:flex;justify-content:space-between;font-size:13px;margin-bottom:3px;">
#             <span style="color:#1A3C6E;font-weight:600;">{svc.replace('_',' ')}</span>
#             <span style="color:#595959;">{count} downloads</span>
#           </div>
#           <div style="background:#e8edf5;border-radius:4px;height:8px;">
#             <div style="background:#2E75B6;width:{pct}%;height:8px;border-radius:4px;"></div>
#           </div>
#         </div>"""

#     html = f"""<!DOCTYPE html>
# <html lang="en">
# <head>
#   <meta charset="UTF-8"/>
#   <meta name="viewport" content="width=device-width, initial-scale=1.0"/>
#   <title>SecureITLab — Download Dashboard</title>
#   <style>
#     * {{ box-sizing: border-box; margin: 0; padding: 0; }}
#     body {{
#       font-family: 'Segoe UI', Arial, sans-serif;
#       background: #eef2f7;
#       color: #333;
#       min-height: 100vh;
#     }}

#     /* ── Top bar ── */
#     .topbar {{
#       background: linear-gradient(135deg, #1A3C6E 0%, #2E75B6 100%);
#       color: white;
#       padding: 0 32px;
#       height: 64px;
#       display: flex;
#       align-items: center;
#       justify-content: space-between;
#       box-shadow: 0 2px 12px rgba(26,60,110,0.3);
#     }}
#     .topbar-brand {{
#       display: flex;
#       align-items: center;
#       gap: 12px;
#     }}
#     .topbar-brand .logo {{
#       width: 36px; height: 36px;
#       background: rgba(255,255,255,0.15);
#       border-radius: 8px;
#       display: flex; align-items: center; justify-content: center;
#       font-size: 18px;
#     }}
#     .topbar-brand h1 {{
#       font-size: 18px;
#       font-weight: 700;
#       letter-spacing: 0.3px;
#     }}
#     .topbar-brand span {{
#       font-size: 12px;
#       opacity: 0.75;
#       display: block;
#     }}
#     .topbar-right {{
#       display: flex;
#       align-items: center;
#       gap: 16px;
#       font-size: 13px;
#       opacity: 0.9;
#     }}
#     .export-btn {{
#       background: rgba(255,255,255,0.15);
#       border: 1px solid rgba(255,255,255,0.35);
#       color: white;
#       padding: 7px 16px;
#       border-radius: 6px;
#       font-size: 13px;
#       cursor: pointer;
#       text-decoration: none;
#       transition: background 0.2s;
#     }}
#     .export-btn:hover {{ background: rgba(255,255,255,0.25); }}

#     /* ── Main content ── */
#     .content {{ padding: 28px 32px; max-width: 1400px; margin: 0 auto; }}

#     /* ── Stat cards ── */
#     .cards {{
#       display: grid;
#       grid-template-columns: repeat(auto-fit, minmax(180px, 1fr));
#       gap: 18px;
#       margin-bottom: 26px;
#     }}
#     .card {{
#       background: white;
#       border-radius: 12px;
#       padding: 22px 24px;
#       box-shadow: 0 1px 6px rgba(0,0,0,0.07);
#       border-left: 4px solid #2E75B6;
#     }}
#     .card.accent {{ border-left-color: #1A3C6E; }}
#     .card.green  {{ border-left-color: #22a06b; }}
#     .card-label  {{ font-size: 11px; text-transform: uppercase; letter-spacing: 0.8px; color: #888; margin-bottom: 8px; }}
#     .card-value  {{ font-size: 32px; font-weight: 700; color: #1A3C6E; line-height: 1; }}
#     .card-sub    {{ font-size: 12px; color: #aaa; margin-top: 5px; }}

#     /* ── Two-col layout ── */
#     .two-col {{ display: grid; grid-template-columns: 1fr 340px; gap: 18px; margin-bottom: 26px; }}
#     @media (max-width: 900px) {{ .two-col {{ grid-template-columns: 1fr; }} }}

#     /* ── Panel (white box) ── */
#     .panel {{
#       background: white;
#       border-radius: 12px;
#       box-shadow: 0 1px 6px rgba(0,0,0,0.07);
#       overflow: hidden;
#     }}
#     .panel-header {{
#       padding: 16px 22px;
#       border-bottom: 1px solid #eef0f5;
#       display: flex;
#       align-items: center;
#       justify-content: space-between;
#     }}
#     .panel-header h2 {{
#       font-size: 14px;
#       font-weight: 700;
#       color: #1A3C6E;
#     }}
#     .panel-header .count-pill {{
#       background: #eef2f7;
#       color: #2E75B6;
#       font-size: 12px;
#       font-weight: 600;
#       padding: 3px 10px;
#       border-radius: 20px;
#     }}
#     .panel-body {{ padding: 20px 22px; }}

#     /* ── Table ── */
#     .tbl-wrap {{
#       overflow-x: auto;
#     }}
#     table {{
#       width: 100%;
#       border-collapse: collapse;
#       font-size: 13px;
#     }}
#     th {{
#       background: #1A3C6E;
#       color: white;
#       padding: 11px 14px;
#       text-align: left;
#       font-size: 11px;
#       font-weight: 600;
#       text-transform: uppercase;
#       letter-spacing: 0.5px;
#       white-space: nowrap;
#     }}
#     td {{
#       padding: 10px 14px;
#       border-bottom: 1px solid #f0f2f7;
#       vertical-align: middle;
#       color: #333;
#     }}
#     tr:last-child td {{ border-bottom: none; }}
#     tr:hover td {{ background: #f0f5ff !important; }}

#     .badge {{
#       background: #eef2f7;
#       color: #1A3C6E;
#       font-size: 11px;
#       font-weight: 600;
#       padding: 3px 9px;
#       border-radius: 20px;
#       white-space: nowrap;
#     }}

#     /* ── Empty state ── */
#     .empty {{
#       text-align: center;
#       padding: 48px 20px;
#       color: #aaa;
#       font-size: 14px;
#     }}
#     .empty .icon {{ font-size: 40px; margin-bottom: 12px; }}

#     /* ── Search bar ── */
#     .search-wrap {{
#       padding: 14px 22px;
#       border-bottom: 1px solid #eef0f5;
#     }}
#     .search-wrap input {{
#       width: 100%;
#       padding: 9px 14px;
#       border: 1px solid #dde2ee;
#       border-radius: 8px;
#       font-size: 13px;
#       outline: none;
#       color: #333;
#     }}
#     .search-wrap input:focus {{ border-color: #2E75B6; box-shadow: 0 0 0 3px rgba(46,117,182,0.1); }}

#     /* ── Footer ── */
#     .footer {{
#       text-align: center;
#       padding: 24px;
#       font-size: 12px;
#       color: #aaa;
#     }}
#   </style>
# </head>
# <body>

# <!-- Top bar -->
# <div class="topbar">
#   <div class="topbar-brand">
#     <div class="logo">🛡️</div>
#     <div>
#       <h1>SecureITLab</h1>
#       <span>Proposal Download Dashboard</span>
#     </div>
#   </div>
#   <div class="topbar-right">
#     <span>👤 {admin}</span>
#     <a class="export-btn" href="/dashboard/export">⬇ Export CSV</a>
#   </div>
# </div>

# <!-- Main content -->
# <div class="content">

#   <!-- Stat cards -->
#   <div class="cards">
#     <div class="card">
#       <div class="card-label">Total Downloads</div>
#       <div class="card-value">{total}</div>
#       <div class="card-sub">All time</div>
#     </div>
#     <div class="card accent">
#       <div class="card-label">Today</div>
#       <div class="card-value">{today_count}</div>
#       <div class="card-sub">{today_str}</div>
#     </div>
#     <div class="card green">
#       <div class="card-label">Unique Services</div>
#       <div class="card-value">{len(service_counts)}</div>
#       <div class="card-sub">Used so far</div>
#     </div>
#     <div class="card">
#       <div class="card-label">Unique Orgs</div>
#       <div class="card-value">{len(set(r.get('Organisation','') for r in rows))}</div>
#       <div class="card-sub">Clients served</div>
#     </div>
#   </div>

#   <!-- Two-col: table + top services -->
#   <div class="two-col">

#     <!-- Download log table -->
#     <div class="panel">
#       <div class="panel-header">
#         <h2>📋 Download Log</h2>
#         <span class="count-pill">{total} records</span>
#       </div>
#       <div class="search-wrap">
#         <input type="text" id="searchInput" placeholder="🔍  Search by client, organisation, service, IP…" oninput="filterTable()"/>
#       </div>
#       <div class="tbl-wrap">
#         {"<div class='empty'><div class='icon'>📭</div>No downloads yet. Generate your first proposal to see it here.</div>" if not rows else f"""
#         <table id="logTable">
#           <thead>
#             <tr>
#               <th>Timestamp</th>
#               <th>Client Name</th>
#               <th>Organisation</th>
#               <th>Service</th>
#               <th>Reference</th>
#               <th>IP Address</th>
#               <th>Filename</th>
#             </tr>
#           </thead>
#           <tbody id="tableBody">
#             {table_rows_html}
#           </tbody>
#         </table>"""}
#       </div>
#     </div>

#     <!-- Top services panel -->
#     <div class="panel">
#       <div class="panel-header">
#         <h2>📊 Top Services</h2>
#       </div>
#       <div class="panel-body">
#         {"<div class='empty'><div class='icon'>📊</div>No data yet.</div>" if not top_html else top_html}
#       </div>
#     </div>

#   </div>
# </div>

# <div class="footer">
#   SecureITLab — Internal Admin Dashboard &nbsp;|&nbsp; Auto-refreshes every 60 seconds
# </div>

# <script>
#   // Search / filter table
#   function filterTable() {{
#     const q = document.getElementById('searchInput').value.toLowerCase();
#     const rows = document.getElementById('tableBody').querySelectorAll('tr');
#     rows.forEach(row => {{
#       row.style.display = row.innerText.toLowerCase().includes(q) ? '' : 'none';
#     }});
#   }}

#   // Auto-refresh every 60 seconds
#   setTimeout(() => location.reload(), 60000);
# </script>
# </body>
# </html>"""

#     return HTMLResponse(content=html)

# # ── Endpoint: export CSV ──────────────────────────────────────────────────────
# @app.get("/dashboard/export")
# def export_csv(admin: str = Depends(_verify_admin)):
#     _ensure_log()
#     content = LOG_FILE.read_text(encoding="utf-8")
#     ts = datetime.now().strftime("%Y%m%d_%H%M%S")
#     return StreamingResponse(
#         io.BytesIO(content.encode("utf-8")),
#         media_type="text/csv",
#         headers={"Content-Disposition": f'attachment; filename="SecureITLab_Downloads_{ts}.csv"'},
#     )

# # ── Document builder ──────────────────────────────────────────────────────────
# def build_docx(req: ProposalRequest, proposal_data: dict, timeline: list) -> io.BytesIO:
#     doc = Document()

#     # ── Page setup (A4, 2.5 cm margins) ──────────────────────────────────────
#     for section in doc.sections:
#         section.page_width  = Cm(21)
#         section.page_height = Cm(29.7)
#         section.top_margin = section.bottom_margin = Cm(2.5)
#         section.left_margin = section.right_margin = Cm(2.5)

#     # ── Document-wide default font ─────────────────────────────────────────
#     style = doc.styles["Normal"]
#     style.font.name = "Arial"
#     style.font.size = Pt(10)

#     # ── Header & footer ────────────────────────────────────────────────────
#     hdr = doc.sections[0].header
#     hdr_para = hdr.paragraphs[0]
#     hdr_para.clear()
#     _add_run(hdr_para, f"SecureITLab — Confidential     |     {req.service_name}",
#              size=8, colour=GREY)
#     _bottom_border(hdr_para, BLUE_HEX, 6)

#     ftr = doc.sections[0].footer
#     ftr_para = ftr.paragraphs[0]
#     ftr_para.clear()
#     _add_run(ftr_para, f"www.secureitlab.com     |     {req.org_name}     |     {req.prop_date}",
#              size=8, colour=GREY)
#     _top_border(ftr_para, CC_HEX, 4)

#     # ══════════════════════════════════════════════════════════════════
#     # COVER PAGE
#     # ══════════════════════════════════════════════════════════════════
#     _spacer(doc, 6)
#     p = doc.add_paragraph()
#     p.alignment = WD_ALIGN_PARAGRAPH.CENTER
#     _add_run(p, "SecureITLab", size=28, bold=True, colour=NAVY)

#     p = doc.add_paragraph()
#     p.alignment = WD_ALIGN_PARAGRAPH.CENTER
#     _add_run(p, "Cybersecurity  |  Data Privacy  |  Data Governance", size=11, colour=BLUE)
#     _bottom_border(p, BLUE_HEX, 6)

#     _spacer(doc, 3)
#     p = doc.add_paragraph()
#     p.alignment = WD_ALIGN_PARAGRAPH.CENTER
#     _add_run(p, "SERVICE PROPOSAL", size=14, bold=True, colour=GREY)

#     p = doc.add_paragraph()
#     p.alignment = WD_ALIGN_PARAGRAPH.CENTER
#     _add_run(p, req.service_name, size=18, bold=True, colour=NAVY)
#     _bottom_border(p, BLUE_HEX, 6)

#     _spacer(doc, 2)
#     p = doc.add_paragraph()
#     p.alignment = WD_ALIGN_PARAGRAPH.CENTER
#     _add_run(p, "Prepared for:", size=11, colour=GREY)

#     p = doc.add_paragraph()
#     p.alignment = WD_ALIGN_PARAGRAPH.CENTER
#     _add_run(p, req.org_name, size=16, bold=True, colour=NAVY)

#     _spacer(doc, 1)
#     p = doc.add_paragraph()
#     p.alignment = WD_ALIGN_PARAGRAPH.CENTER
#     _add_run(p, f"Proposal Date: {req.prop_date}", size=10, colour=GREY)

#     p = doc.add_paragraph()
#     p.alignment = WD_ALIGN_PARAGRAPH.CENTER
#     _add_run(p, f"Reference: {req.prop_ref}", size=10, colour=GREY)

#     doc.add_page_break()

#     # ══════════════════════════════════════════════════════════════════
#     # S1  Organisation Overview
#     # ══════════════════════════════════════════════════════════════════
#     _heading(doc, "1.  Organisation Overview")
#     _rule(doc)

#     fields = [
#         ("Organisation Name", req.org_name),
#         ("Industry / Sector",  req.org_industry),
#         ("Registered Address", req.org_address),
#         ("Primary Contact",    req.org_contact),
#         ("Email",              req.org_email),
#         ("Phone",              req.org_phone),
#         ("Organisation Background", req.org_bg),
#         ("Engagement Objective",    req.org_obj),
#     ]
#     for label, val in fields:
#         p = doc.add_paragraph()
#         _add_run(p, f"{label}:  ", bold=True, colour=NAVY)
#         _add_run(p, val or "—", colour=GREY)

#     doc.add_page_break()

#     # ══════════════════════════════════════════════════════════════════
#     # S2  About SecureITLab  (boilerplate)
#     # ══════════════════════════════════════════════════════════════════
#     _heading(doc, "2.  About SecureITLab")
#     _rule(doc)

#     _body(doc, ("SecureITLab is a remote-first cybersecurity, data privacy, and data governance "
#                 "consulting firm with over 20 years of industry experience. Headquartered in the "
#                 "Kingdom of Bahrain, with offices in the Philippines and India, SecureITLab serves "
#                 "clients across the Middle East, Asia-Pacific, and beyond."))
#     _body(doc, ("Our team of certified professionals (CISSP, CISM, CISA, CEH, ISO 27001 LA/LI, CDPSE) "
#                 "delivers tailored, risk-driven solutions that empower organisations to navigate the "
#                 "complex landscape of cybersecurity and compliance with confidence."))
#     _body(doc, "SecureITLab is certified to ISO 27001:2022, ISO 27701:2019, and ISO 22301:2019.")

#     _subheading(doc, "Core Capabilities")
#     caps = [
#         "Consulting — Cybersecurity and Data Privacy",
#         "Assurance — Penetration Testing, Vulnerability Assessment and Red Teaming",
#         "Audit and Compliance — IT Audit, Cloud Audit and ISO Compliance",
#         "Services Augmentation — vCISO, vCDO and Expert Resourcing",
#         "Managed Services — SOC, NOC, Vulnerability Management and Take-Down",
#     ]
#     for c in caps:
#         _bullet(doc, c)

#     doc.add_page_break()

#     # ══════════════════════════════════════════════════════════════════
#     # S3  Service Description
#     # ══════════════════════════════════════════════════════════════════
#     _heading(doc, "3.  Service Description")
#     _rule(doc)

#     p = doc.add_paragraph()
#     _add_run(p, req.service_name, size=12, bold=True, colour=BLUE)
#     _body(doc, req.service_desc)

#     _subheading(doc, "Key Deliverables")
#     for d in req.service_deliverables:
#         _bullet(doc, d)

#     _subheading(doc, "Methodology")
#     _body(doc, ("SecureITLab employs industry-recognised methodologies, standards, and frameworks "
#                 "throughout this engagement. All activities are conducted in alignment with the agreed "
#                 "scope and subject to the client's change-control and information security policies."))
#     for phase_text in proposal_data["methodology"]:
#         _body(doc, phase_text)

#     doc.add_page_break()

#     # ══════════════════════════════════════════════════════════════════
#     # S4  Scope of Work
#     # ══════════════════════════════════════════════════════════════════
#     _heading(doc, "4.  Scope of Work")
#     _rule(doc)
#     _body(doc, "Define the precise boundaries of the engagement. Items not listed are considered out of scope.")

#     _subheading(doc, "In Scope")
#     _bullet_list(doc, req.inscope, fallback="To be defined.")

#     _subheading(doc, "Out of Scope")
#     _bullet_list(doc, req.outscope, fallback="To be defined.")

#     _subheading(doc, "Assumptions and Dependencies")
#     _bullet_list(doc, req.assumptions, fallback="To be confirmed.")

#     _subheading(doc, "Client Responsibilities")
#     _bullet_list(doc, req.responsibilities, fallback="To be confirmed.")

#     doc.add_page_break()

#     # ══════════════════════════════════════════════════════════════════
#     # S5  Project Timeline
#     # ══════════════════════════════════════════════════════════════════
#     _heading(doc, "5.  Project Timeline")
#     _rule(doc)
#     _body(doc, "The indicative timeline below is subject to confirmation upon contract award.")

#     p = doc.add_paragraph()
#     _add_run(p, "Proposed Start: ", bold=True, colour=NAVY)
#     _add_run(p, req.start_date or "TBD", colour=GREY)
#     _add_run(p, "     Proposed End: ", bold=True, colour=NAVY)
#     _add_run(p, req.end_date or "TBD", colour=GREY)

#     col_w = [Cm(3.2), Cm(6.2), Cm(2.8), Cm(5.5)]
#     tbl = doc.add_table(rows=1, cols=4)
#     tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
#     _set_col_widths(tbl, col_w)

#     hdr_cells = tbl.rows[0].cells
#     for i, h in enumerate(["Phase", "Activities", "Duration", "Deliverables"]):
#         _header_cell(hdr_cells[i], h, col_w[i])

#     for idx, row in enumerate(timeline):
#         cells = tbl.add_row().cells
#         fill = LBLUE_HEX if idx % 2 == 1 else WHITE_HEX
#         _data_cell(cells[0], row.phase,        col_w[0], fill)
#         _data_cell(cells[1], row.activities,   col_w[1], fill)
#         _data_cell(cells[2], row.duration,     col_w[2], fill)
#         _data_cell(cells[3], row.deliverables, col_w[3], fill)

#     doc.add_page_break()

#     # ══════════════════════════════════════════════════════════════════
#     # S6  Investment Summary
#     # ══════════════════════════════════════════════════════════════════
#     _heading(doc, "6.  Investment Summary")
#     _rule(doc)
#     _body(doc, "All fees are exclusive of applicable taxes. Payment terms are as per the Master Service Agreement or as agreed in writing.")

#     p = doc.add_paragraph()
#     _add_run(p, f"Currency: {req.currency}", bold=True, colour=NAVY)

#     cost_rows = req.cost_rows
#     subtotal  = sum(r.total for r in cost_rows)
#     tax_amt   = round(subtotal * req.tax_rate / 100, 2)
#     grand     = round(subtotal + tax_amt, 2)

#     ccol = [Cm(4.2), Cm(5.4), Cm(2.8), Cm(1.6), Cm(3.7)]
#     ctbl = doc.add_table(rows=1, cols=5)
#     ctbl.alignment = WD_TABLE_ALIGNMENT.LEFT
#     _set_col_widths(ctbl, ccol)

#     ch = ctbl.rows[0].cells
#     for i, h in enumerate(["Line Item", "Description", "Unit Price", "Qty", f"Total ({req.currency})"]):
#         _header_cell(ch[i], h, ccol[i])

#     for idx, r in enumerate(cost_rows):
#         cells = ctbl.add_row().cells
#         fill  = LBLUE_HEX if idx % 2 == 1 else WHITE_HEX
#         _data_cell(cells[0], r.item,             ccol[0], fill)
#         _data_cell(cells[1], r.desc,             ccol[1], fill)
#         _data_cell(cells[2], _fmt(r.unit_price), ccol[2], fill)
#         _data_cell(cells[3], str(r.qty),         ccol[3], fill)
#         _data_cell(cells[4], _fmt(r.total),      ccol[4], fill)

#     _summary_row(ctbl, ccol, f"Subtotal",              _fmt(subtotal), LBLUE_HEX, NAVY_HEX)
#     _summary_row(ctbl, ccol, f"Tax ({req.tax_rate}%)", _fmt(tax_amt),  LBLUE_HEX, NAVY_HEX)
#     _summary_row(ctbl, ccol, "TOTAL INVESTMENT",
#                  f"{req.currency} {_fmt(grand)}", NAVY_HEX, WHITE_HEX, bold_label=True)

#     _subheading(doc, "Payment Schedule")
#     _bullet_list(doc, req.payments, fallback="As agreed.")
#     _body(doc, "Validity: This proposal is valid for 30 days from the date of issue.")

#     doc.add_page_break()

#     # ══════════════════════════════════════════════════════════════════
#     # S7  Terms and Conditions  (boilerplate)
#     # ══════════════════════════════════════════════════════════════════
#     _heading(doc, "7.  Terms and Conditions")
#     _rule(doc)
#     _body(doc, ("This proposal is subject to SecureITLab's standard Master Service Agreement (MSA). "
#                 "Key terms are summarised below; the MSA governs in the event of any conflict."))

#     tc_items = [
#         ("Confidentiality",
#          "Both parties agree to maintain strict confidentiality of all information exchanged. "
#          "SecureITLab will not disclose client data to third parties without prior written consent."),
#         ("Intellectual Property",
#          "All deliverables become the property of the client upon receipt of full payment. "
#          "SecureITLab retains rights to its pre-existing methodologies, tools, and frameworks."),
#         ("Limitation of Liability",
#          "SecureITLab's aggregate liability shall not exceed the total fees paid under this engagement. "
#          "SecureITLab is not liable for indirect, consequential, or incidental damages."),
#         ("Governing Law",
#          "This proposal and any resulting engagement shall be governed by the laws of the Kingdom of "
#          "Bahrain, unless otherwise agreed in writing."),
#         ("Amendments",
#          "Any change in scope, timeline, or cost requires a written Change Request signed by both "
#          "parties prior to implementation."),
#     ]
#     for title, body in tc_items:
#         p = doc.add_paragraph()
#         _add_run(p, f"{title}:  ", bold=True, colour=NAVY)
#         _add_run(p, body, colour=GREY)

#     doc.add_page_break()

#     # ══════════════════════════════════════════════════════════════════
#     # S8  Acceptance and Signatures
#     # ══════════════════════════════════════════════════════════════════
#     _heading(doc, "8.  Acceptance and Signatures")
#     _rule(doc)
#     _body(doc, ("By signing below, both parties confirm their agreement to the terms set out in this "
#                 "proposal and authorise SecureITLab to commence the engagement as described."))

#     _spacer(doc, 2)

#     sig_col = [Cm(7.5), Cm(2.0), Cm(7.5)]
#     stbl = doc.add_table(rows=4, cols=3)
#     _set_col_widths(stbl, sig_col)

#     header_row = stbl.rows[0]
#     _sig_header_cell(header_row.cells[0], f"Authorised Signatory — {req.org_name}", sig_col[0])
#     _sig_header_cell(header_row.cells[2], "Authorised Signatory — SecureITLab",     sig_col[2])

#     for row in stbl.rows[1:]:
#         for cell in row.cells:
#             cell.text = ""

#     _spacer(doc, 2)
#     _rule(doc)

#     p = doc.add_paragraph()
#     p.alignment = WD_ALIGN_PARAGRAPH.CENTER
#     _add_run(p, "SecureITLab  |  info@secureitlab.com  |  www.secureitlab.com", size=9, colour=GREY)

#     p = doc.add_paragraph()
#     p.alignment = WD_ALIGN_PARAGRAPH.CENTER
#     _add_run(p, "Bahrain  |  Philippines  |  India", size=9, colour=GREY)

#     buf = io.BytesIO()
#     doc.save(buf)
#     buf.seek(0)
#     return buf


# # ── Low-level helpers ─────────────────────────────────────────────────────────

# def _add_run(para, text, size=10, bold=False, italic=False, colour=None):
#     run = para.add_run(text)
#     run.font.name = "Arial"
#     run.font.size = Pt(size)
#     run.font.bold = bold
#     run.font.italic = italic
#     if colour:
#         run.font.color.rgb = colour
#     return run

# def _heading(doc, text):
#     p = doc.add_paragraph()
#     p.paragraph_format.space_before = Pt(14)
#     p.paragraph_format.space_after  = Pt(6)
#     _add_run(p, text, size=14, bold=True, colour=NAVY)

# def _subheading(doc, text):
#     p = doc.add_paragraph()
#     p.paragraph_format.space_before = Pt(8)
#     p.paragraph_format.space_after  = Pt(4)
#     _add_run(p, text, size=11, bold=True, colour=NAVY)

# def _body(doc, text):
#     p = doc.add_paragraph()
#     p.paragraph_format.space_before = Pt(2)
#     p.paragraph_format.space_after  = Pt(4)
#     _add_run(p, text, colour=GREY)

# def _bullet(doc, text):
#     p = doc.add_paragraph(style="List Bullet")
#     _add_run(p, text, colour=GREY)

# def _bullet_list(doc, items, fallback=""):
#     clean = [i for i in items if i.strip()]
#     if clean:
#         for item in clean:
#             _bullet(doc, item)
#     elif fallback:
#         _body(doc, fallback)

# def _spacer(doc, lines=1):
#     for _ in range(lines):
#         doc.add_paragraph()

# def _rule(doc):
#     p = doc.add_paragraph()
#     _bottom_border(p, BLUE_HEX, 6)

# def _fmt(val):
#     return f"{val:,.2f}"

# def _bottom_border(para, colour_hex, size):
#     pPr = para._p.get_or_add_pPr()
#     pBdr = OxmlElement("w:pBdr")
#     bottom = OxmlElement("w:bottom")
#     bottom.set(qn("w:val"), "single")
#     bottom.set(qn("w:sz"), str(size))
#     bottom.set(qn("w:space"), "1")
#     bottom.set(qn("w:color"), colour_hex)
#     pBdr.append(bottom)
#     pPr.append(pBdr)

# def _top_border(para, colour_hex, size):
#     pPr = para._p.get_or_add_pPr()
#     pBdr = OxmlElement("w:pBdr")
#     top = OxmlElement("w:top")
#     top.set(qn("w:val"), "single")
#     top.set(qn("w:sz"), str(size))
#     top.set(qn("w:space"), "1")
#     top.set(qn("w:color"), colour_hex)
#     pBdr.append(top)
#     pPr.append(pBdr)

# def _set_col_widths(table, widths):
#     tbl_el = table._tbl
#     tblPr  = tbl_el.find(qn("w:tblPr")) or OxmlElement("w:tblPr")
#     tblW   = OxmlElement("w:tblW")
#     total  = sum(int(w.cm * 914400 / 360) for w in widths)
#     tblW.set(qn("w:w"),    str(total))
#     tblW.set(qn("w:type"), "dxa")
#     tblPr.append(tblW)

# def _twips(cm_val):
#     return int(cm_val.cm * 567)

# def _shade_cell(cell, fill_hex):
#     tcPr = cell._tc.get_or_add_tcPr()
#     shd  = OxmlElement("w:shd")
#     shd.set(qn("w:val"),   "clear")
#     shd.set(qn("w:color"), "auto")
#     shd.set(qn("w:fill"),  fill_hex)
#     tcPr.append(shd)

# def _cell_borders(cell, colour_hex, size):
#     tcPr = cell._tc.get_or_add_tcPr()
#     tcBorders = OxmlElement("w:tcBorders")
#     for side in ("top", "left", "bottom", "right"):
#         el = OxmlElement(f"w:{side}")
#         el.set(qn("w:val"),   "single")
#         el.set(qn("w:sz"),    str(size))
#         el.set(qn("w:color"), colour_hex)
#         tcBorders.append(el)
#     tcPr.append(tcBorders)

# def _cell_width(cell, cm_val):
#     tcPr = cell._tc.get_or_add_tcPr()
#     tcW  = OxmlElement("w:tcW")
#     tcW.set(qn("w:w"),    str(_twips(cm_val)))
#     tcW.set(qn("w:type"), "dxa")
#     tcPr.append(tcW)

# def _header_cell(cell, text, width):
#     cell.text = ""
#     _shade_cell(cell, NAVY_HEX)
#     _cell_borders(cell, BLUE_HEX, 6)
#     _cell_width(cell, width)
#     p = cell.paragraphs[0]
#     p.alignment = WD_ALIGN_PARAGRAPH.CENTER
#     _add_run(p, text, size=9, bold=True, colour=WHITE)

# def _data_cell(cell, text, width, fill=WHITE_HEX):
#     cell.text = ""
#     _shade_cell(cell, fill)
#     _cell_borders(cell, CC_HEX, 4)
#     _cell_width(cell, width)
#     p = cell.paragraphs[0]
#     _add_run(p, text or "", size=9, colour=GREY)

# def _summary_row(table, col_widths, label_text, value_text, label_fill, label_colour_hex, bold_label=False):
#     row   = table.add_row()
#     cells = row.cells
#     merged = cells[0].merge(cells[1]).merge(cells[2]).merge(cells[3])
#     _shade_cell(merged, label_fill)
#     _cell_borders(merged, BLUE_HEX, 6)
#     p = merged.paragraphs[0]
#     p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
#     lc = WHITE if label_fill == NAVY_HEX else NAVY
#     _add_run(p, label_text, size=10, bold=bold_label, colour=lc)

#     vc = cells[4]
#     _shade_cell(vc, label_fill)
#     _cell_borders(vc, BLUE_HEX, 6)
#     _cell_width(vc, col_widths[4])
#     vp = vc.paragraphs[0]
#     vc_text_colour = WHITE if label_fill == NAVY_HEX else GREY
#     _add_run(vp, value_text, size=10, bold=bold_label, colour=vc_text_colour)

# def _sig_header_cell(cell, text, width):
#     _cell_width(cell, width)
#     tcPr = cell._tc.get_or_add_tcPr()
#     tcBorders = OxmlElement("w:tcBorders")
#     bottom = OxmlElement("w:bottom")
#     bottom.set(qn("w:val"),   "single")
#     bottom.set(qn("w:sz"),    "6")
#     bottom.set(qn("w:color"), NAVY_HEX)
#     tcBorders.append(bottom)
#     tcPr.append(tcBorders)
#     p = cell.paragraphs[0]
#     p.paragraph_format.space_before = Pt(40)
#     _add_run(p, text, size=9, bold=True, colour=NAVY)





















"""
main.py  —  SecureITLab Proposal Generator Backend
Run with:  uvicorn main:app --reload --port 8000

NEW in this version:
  • Every /generate call is logged to downloads.csv
  • Admin dashboard at /dashboard  (login required)
  • Admin credentials set via ADMIN_USER / ADMIN_PASS env vars
    (defaults: admin / secureitlab2024)
"""

from fastapi import FastAPI, HTTPException, Request, Depends
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse, JSONResponse, HTMLResponse, RedirectResponse
from fastapi.security import HTTPBasic, HTTPBasicCredentials
from pydantic import BaseModel
from typing import List, Optional
import io, os, csv, secrets
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from lxml import etree

from proposal_data import PROPOSALS

# ── Admin credentials (set env vars to override) ─────────────────────────────
ADMIN_USER = os.getenv("ADMIN_USER", "admin")
ADMIN_PASS = os.getenv("ADMIN_PASS", "proposal_sitl_2026")

# ── CSV log file ──────────────────────────────────────────────────────────────
LOG_FILE = Path("downloads.csv")
CSV_HEADERS = ["Timestamp", "Client Name", "Organisation", "Service", "Reference", "IP Address", "Filename"]

def _ensure_log():
    """Create the CSV with headers if it doesn't exist."""
    if not LOG_FILE.exists():
        with open(LOG_FILE, "w", newline="", encoding="utf-8") as f:
            writer = csv.writer(f)
            writer.writerow(CSV_HEADERS)

def _log_download(req_data, ip: str, filename: str):
    """Append one row to downloads.csv."""
    _ensure_log()
    with open(LOG_FILE, "a", newline="", encoding="utf-8") as f:
        writer = csv.writer(f)
        writer.writerow([
            datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            req_data.org_contact or "—",
            req_data.org_name    or "—",
            req_data.service_name,
            req_data.prop_ref    or "—",
            ip,
            filename,
        ])

def _read_log() -> list:
    """Return all rows from the CSV as a list of dicts."""
    _ensure_log()
    with open(LOG_FILE, "r", encoding="utf-8") as f:
        return list(csv.DictReader(f))

# ── Colour helpers ────────────────────────────────────────────────────────────
NAVY   = RGBColor(0x1A, 0x3C, 0x6E)
BLUE   = RGBColor(0x2E, 0x75, 0xB6)
LBLUE  = RGBColor(0xD5, 0xE8, 0xF0)
GREY   = RGBColor(0x59, 0x59, 0x59)
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)

NAVY_HEX  = "1A3C6E"
BLUE_HEX  = "2E75B6"
LBLUE_HEX = "D5E8F0"
GREY_HEX  = "595959"
WHITE_HEX = "FFFFFF"
CC_HEX    = "CCCCCC"

# ── Pydantic models ───────────────────────────────────────────────────────────
class CostRow(BaseModel):
    item: str = ""
    desc: str = ""
    unit_price: float = 0
    qty: int = 1
    total: float = 0

class TimelineRow(BaseModel):
    phase: str = ""
    activities: str = ""
    duration: str = ""
    deliverables: str = ""

class ProposalRequest(BaseModel):
    # Service
    service_key: str
    service_name: str
    service_desc: str
    service_deliverables: List[str] = []

    # Organisation
    org_name: str = ""
    org_industry: str = ""
    org_address: str = ""
    org_contact: str = ""
    org_email: str = ""
    org_phone: str = ""
    org_bg: str = ""
    org_obj: str = ""

    # Proposal meta
    prop_date: str = ""
    prop_ref: str = ""
    start_date: str = ""
    end_date: str = ""
    currency: str = "USD"
    tax_rate: float = 0

    # Scope lists
    inscope: List[str] = []
    outscope: List[str] = []
    assumptions: List[str] = []
    responsibilities: List[str] = []
    payments: List[str] = []

    # Tables
    timeline_rows: List[TimelineRow] = []
    cost_rows: List[CostRow] = []

# ── FastAPI app ───────────────────────────────────────────────────────────────
app = FastAPI(title="SecureITLab Proposal Generator", version="2.0.0")
security = HTTPBasic()

app.add_middleware(
    CORSMiddleware,
    # allow_origins=["*"],
    allow_origins=["http://38.242.247.151:8015"],
    allow_methods=["*"],
    allow_headers=["*"],
)

# Serve the HTML front-end at /
if os.path.exists("SecureITLab_Proposal_Generator.html"):
    from fastapi.responses import FileResponse
    @app.get("/")
    def serve_frontend():
        return FileResponse("SecureITLab_Proposal_Generator.html")

# ── Admin auth helper ─────────────────────────────────────────────────────────
def _verify_admin(credentials: HTTPBasicCredentials = Depends(security)):
    ok_user = secrets.compare_digest(credentials.username.encode(), ADMIN_USER.encode())
    ok_pass = secrets.compare_digest(credentials.password.encode(), ADMIN_PASS.encode())
    if not (ok_user and ok_pass):
        raise HTTPException(
            status_code=401,
            detail="Invalid credentials",
            headers={"WWW-Authenticate": "Basic"},
        )
    return credentials.username

# ── Endpoint: list services ───────────────────────────────────────────────────
@app.get("/services")
def list_services():
    return {k: {"phases": len(v["timeline"])} for k, v in PROPOSALS.items()}

# ── Endpoint: get service detail ──────────────────────────────────────────────
@app.get("/services/{key}")
def get_service(key: str):
    if key not in PROPOSALS:
        raise HTTPException(404, f"Service key '{key}' not found")
    data = PROPOSALS[key]
    return {
        "methodology": data["methodology"],
        "timeline": [
            {"phase": t[0], "activities": t[1], "duration": t[2], "deliverables": t[3]}
            for t in data["timeline"]
        ]
    }

# ── Endpoint: generate proposal docx ─────────────────────────────────────────
@app.post("/generate")
def generate_proposal(req: ProposalRequest, request: Request):
    if req.service_key not in PROPOSALS:
        raise HTTPException(404, f"Service key '{req.service_key}' not found")

    proposal_data = PROPOSALS[req.service_key]

    if not req.timeline_rows:
        timeline = [
            TimelineRow(phase=t[0], activities=t[1], duration=t[2], deliverables=t[3])
            for t in proposal_data["timeline"]
        ]
    else:
        timeline = req.timeline_rows

    buf = build_docx(req, proposal_data, timeline)

    safe_name = (req.org_name or "Proposal").replace(" ", "_")
    safe_svc  = req.service_name.replace(" ", "_").replace("/", "-")
    filename  = f"Proposal_{safe_name}_{safe_svc}.docx"

    # ── Log the download ──────────────────────────────────────────────────────
    client_ip = request.headers.get("X-Forwarded-For", request.client.host)
    _log_download(req, client_ip, filename)

    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )

# ── Endpoint: admin dashboard ─────────────────────────────────────────────────
@app.get("/dashboard", response_class=HTMLResponse)
def dashboard(admin: str = Depends(_verify_admin)):
    rows = _read_log()
    total      = len(rows)
    today_str  = datetime.now().strftime("%Y-%m-%d")
    today_count = sum(1 for r in rows if r["Timestamp"].startswith(today_str))

    # Count by service
    service_counts: dict = {}
    for r in rows:
        svc = r.get("Service", "Unknown")
        service_counts[svc] = service_counts.get(svc, 0) + 1
    top_services = sorted(service_counts.items(), key=lambda x: x[1], reverse=True)[:5]

    # Build table rows HTML (newest first)
    table_rows_html = ""
    for i, r in enumerate(reversed(rows)):
        bg = "#f8faff" if i % 2 == 0 else "#ffffff"
        table_rows_html += f"""
        <tr style="background:{bg};">
          <td>{r.get('Timestamp','—')}</td>
          <td>{r.get('Client Name','—')}</td>
          <td>{r.get('Organisation','—')}</td>
          <td><span class="badge">{r.get('Service','—')}</span></td>
          <td>{r.get('Reference','—')}</td>
          <td>{r.get('IP Address','—')}</td>
          <td style="font-size:11px;color:#888;">{r.get('Filename','—')}</td>
        </tr>"""

    # Build top services HTML
    top_html = ""
    for svc, count in top_services:
        pct = round(count / total * 100) if total else 0
        svc_display = svc.replace('_', ' ')
        top_html += (
            '<div style="margin-bottom:10px;">'
            '<div style="display:flex;justify-content:space-between;font-size:13px;margin-bottom:3px;">'
            '<span style="color:#1A3C6E;font-weight:600;">' + svc_display + '</span>'
            '<span style="color:#595959;">' + str(count) + ' downloads</span>'
            '</div>'
            '<div style="background:#e8edf5;border-radius:4px;height:8px;">'
            '<div style="background:#2E75B6;width:' + str(pct) + '%;height:8px;border-radius:4px;"></div>'
            '</div>'
            '</div>'
        )

    # Pre-build table section HTML (avoids nested f-string — Python 3.8 compatible)
    if not rows:
        table_section_html = "<div class='empty'><div class='icon'>📭</div>No downloads yet. Generate your first proposal to see it here.</div>"
    else:
        table_section_html = (
            '<table id="logTable">'
            '<thead><tr>'
            '<th>Timestamp</th>'
            '<th>Client Name</th>'
            '<th>Organisation</th>'
            '<th>Service</th>'
            '<th>Reference</th>'
            '<th>IP Address</th>'
            '<th>Filename</th>'
            '</tr></thead>'
            '<tbody id="tableBody">'
            + table_rows_html +
            '</tbody></table>'
        )

    # Pre-build top services section HTML
    if not top_html:
        top_section_html = "<div class='empty'><div class='icon'>📊</div>No data yet.</div>"
    else:
        top_section_html = top_html

    unique_orgs = len(set(r.get('Organisation', '') for r in rows))

    html = f"""<!DOCTYPE html>
<html lang="en">
<head>
  <meta charset="UTF-8"/>
  <meta name="viewport" content="width=device-width, initial-scale=1.0"/>
  <title>SecureITLab — Download Dashboard</title>
  <style>
    * {{ box-sizing: border-box; margin: 0; padding: 0; }}
    body {{
      font-family: 'Segoe UI', Arial, sans-serif;
      background: #eef2f7;
      color: #333;
      min-height: 100vh;
    }}

    /* ── Top bar ── */
    .topbar {{
      background: linear-gradient(135deg, #1A3C6E 0%, #2E75B6 100%);
      color: white;
      padding: 0 32px;
      height: 64px;
      display: flex;
      align-items: center;
      justify-content: space-between;
      box-shadow: 0 2px 12px rgba(26,60,110,0.3);
    }}
    .topbar-brand {{
      display: flex;
      align-items: center;
      gap: 12px;
    }}
    .topbar-brand .logo {{
      width: 36px; height: 36px;
      background: rgba(255,255,255,0.15);
      border-radius: 8px;
      display: flex; align-items: center; justify-content: center;
      font-size: 18px;
    }}
    .topbar-brand h1 {{
      font-size: 18px;
      font-weight: 700;
      letter-spacing: 0.3px;
    }}
    .topbar-brand span {{
      font-size: 12px;
      opacity: 0.75;
      display: block;
    }}
    .topbar-right {{
      display: flex;
      align-items: center;
      gap: 16px;
      font-size: 13px;
      opacity: 0.9;
    }}
    .export-btn {{
      background: rgba(255,255,255,0.15);
      border: 1px solid rgba(255,255,255,0.35);
      color: white;
      padding: 7px 16px;
      border-radius: 6px;
      font-size: 13px;
      cursor: pointer;
      text-decoration: none;
      transition: background 0.2s;
    }}
    .export-btn:hover {{ background: rgba(255,255,255,0.25); }}

    /* ── Main content ── */
    .content {{ padding: 28px 32px; max-width: 1400px; margin: 0 auto; }}

    /* ── Stat cards ── */
    .cards {{
      display: grid;
      grid-template-columns: repeat(auto-fit, minmax(180px, 1fr));
      gap: 18px;
      margin-bottom: 26px;
    }}
    .card {{
      background: white;
      border-radius: 12px;
      padding: 22px 24px;
      box-shadow: 0 1px 6px rgba(0,0,0,0.07);
      border-left: 4px solid #2E75B6;
    }}
    .card.accent {{ border-left-color: #1A3C6E; }}
    .card.green  {{ border-left-color: #22a06b; }}
    .card-label  {{ font-size: 11px; text-transform: uppercase; letter-spacing: 0.8px; color: #888; margin-bottom: 8px; }}
    .card-value  {{ font-size: 32px; font-weight: 700; color: #1A3C6E; line-height: 1; }}
    .card-sub    {{ font-size: 12px; color: #aaa; margin-top: 5px; }}

    /* ── Two-col layout ── */
    .two-col {{ display: grid; grid-template-columns: 1fr 340px; gap: 18px; margin-bottom: 26px; }}
    @media (max-width: 900px) {{ .two-col {{ grid-template-columns: 1fr; }} }}

    /* ── Panel (white box) ── */
    .panel {{
      background: white;
      border-radius: 12px;
      box-shadow: 0 1px 6px rgba(0,0,0,0.07);
      overflow: hidden;
    }}
    .panel-header {{
      padding: 16px 22px;
      border-bottom: 1px solid #eef0f5;
      display: flex;
      align-items: center;
      justify-content: space-between;
    }}
    .panel-header h2 {{
      font-size: 14px;
      font-weight: 700;
      color: #1A3C6E;
    }}
    .panel-header .count-pill {{
      background: #eef2f7;
      color: #2E75B6;
      font-size: 12px;
      font-weight: 600;
      padding: 3px 10px;
      border-radius: 20px;
    }}
    .panel-body {{ padding: 20px 22px; }}

    /* ── Table ── */
    .tbl-wrap {{
      overflow-x: auto;
    }}
    table {{
      width: 100%;
      border-collapse: collapse;
      font-size: 13px;
    }}
    th {{
      background: #1A3C6E;
      color: white;
      padding: 11px 14px;
      text-align: left;
      font-size: 11px;
      font-weight: 600;
      text-transform: uppercase;
      letter-spacing: 0.5px;
      white-space: nowrap;
    }}
    td {{
      padding: 10px 14px;
      border-bottom: 1px solid #f0f2f7;
      vertical-align: middle;
      color: #333;
    }}
    tr:last-child td {{ border-bottom: none; }}
    tr:hover td {{ background: #f0f5ff !important; }}

    .badge {{
      background: #eef2f7;
      color: #1A3C6E;
      font-size: 11px;
      font-weight: 600;
      padding: 3px 9px;
      border-radius: 20px;
      white-space: nowrap;
    }}

    /* ── Empty state ── */
    .empty {{
      text-align: center;
      padding: 48px 20px;
      color: #aaa;
      font-size: 14px;
    }}
    .empty .icon {{ font-size: 40px; margin-bottom: 12px; }}

    /* ── Search bar ── */
    .search-wrap {{
      padding: 14px 22px;
      border-bottom: 1px solid #eef0f5;
    }}
    .search-wrap input {{
      width: 100%;
      padding: 9px 14px;
      border: 1px solid #dde2ee;
      border-radius: 8px;
      font-size: 13px;
      outline: none;
      color: #333;
    }}
    .search-wrap input:focus {{ border-color: #2E75B6; box-shadow: 0 0 0 3px rgba(46,117,182,0.1); }}

    /* ── Footer ── */
    .footer {{
      text-align: center;
      padding: 24px;
      font-size: 12px;
      color: #aaa;
    }}
  </style>
</head>
<body>

<!-- Top bar -->
<div class="topbar">
  <div class="topbar-brand">
    <div class="logo">🛡️</div>
    <div>
      <h1>SecureITLab</h1>
      <span>Proposal Download Dashboard</span>
    </div>
  </div>
  <div class="topbar-right">
    <span>👤 {admin}</span>
    <a class="export-btn" href="/dashboard/export">⬇ Export CSV</a>
  </div>
</div>

<!-- Main content -->
<div class="content">

  <!-- Stat cards -->
  <div class="cards">
    <div class="card">
      <div class="card-label">Total Downloads</div>
      <div class="card-value">{total}</div>
      <div class="card-sub">All time</div>
    </div>
    <div class="card accent">
      <div class="card-label">Today</div>
      <div class="card-value">{today_count}</div>
      <div class="card-sub">{today_str}</div>
    </div>
    <div class="card green">
      <div class="card-label">Unique Services</div>
      <div class="card-value">{len(service_counts)}</div>
      <div class="card-sub">Used so far</div>
    </div>
    <div class="card">
      <div class="card-label">Unique Orgs</div>
      <div class="card-value">{unique_orgs}</div>
      <div class="card-sub">Clients served</div>
    </div>
  </div>

  <!-- Two-col: table + top services -->
  <div class="two-col">

    <!-- Download log table -->
    <div class="panel">
      <div class="panel-header">
        <h2>📋 Download Log</h2>
        <span class="count-pill">{total} records</span>
      </div>
      <div class="search-wrap">
        <input type="text" id="searchInput" placeholder="🔍  Search by client, organisation, service, IP…" oninput="filterTable()"/>
      </div>
      <div class="tbl-wrap">
        {table_section_html}
      </div>
    </div>

    <!-- Top services panel -->
    <div class="panel">
      <div class="panel-header">
        <h2>📊 Top Services</h2>
      </div>
      <div class="panel-body">
        {top_section_html}
      </div>
    </div>

  </div>
</div>

<div class="footer">
  SecureITLab — Internal Admin Dashboard &nbsp;|&nbsp; Auto-refreshes every 60 seconds
</div>

<script>
  // Search / filter table
  function filterTable() {{
    const q = document.getElementById('searchInput').value.toLowerCase();
    const rows = document.getElementById('tableBody').querySelectorAll('tr');
    rows.forEach(row => {{
      row.style.display = row.innerText.toLowerCase().includes(q) ? '' : 'none';
    }});
  }}

  // Auto-refresh every 60 seconds
  setTimeout(() => location.reload(), 60000);
</script>
</body>
</html>"""

    return HTMLResponse(content=html)

# ── Endpoint: export CSV ──────────────────────────────────────────────────────
@app.get("/dashboard/export")
def export_csv(admin: str = Depends(_verify_admin)):
    _ensure_log()
    content = LOG_FILE.read_text(encoding="utf-8")
    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
    return StreamingResponse(
        io.BytesIO(content.encode("utf-8")),
        media_type="text/csv",
        headers={"Content-Disposition": f'attachment; filename="SecureITLab_Downloads_{ts}.csv"'},
    )

# ── Document builder ──────────────────────────────────────────────────────────
def build_docx(req: ProposalRequest, proposal_data: dict, timeline: list) -> io.BytesIO:
    doc = Document()

    # ── Page setup (A4, 2.5 cm margins) ──────────────────────────────────────
    for section in doc.sections:
        section.page_width  = Cm(21)
        section.page_height = Cm(29.7)
        section.top_margin = section.bottom_margin = Cm(2.5)
        section.left_margin = section.right_margin = Cm(2.5)

    # ── Document-wide default font ─────────────────────────────────────────
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(10)

    # ── Header & footer ────────────────────────────────────────────────────
    hdr = doc.sections[0].header
    hdr_para = hdr.paragraphs[0]
    hdr_para.clear()
    _add_run(hdr_para, f"SecureITLab — Confidential     |     {req.service_name}",
             size=8, colour=GREY)
    _bottom_border(hdr_para, BLUE_HEX, 6)

    ftr = doc.sections[0].footer
    ftr_para = ftr.paragraphs[0]
    ftr_para.clear()
    _add_run(ftr_para, f"www.secureitlab.com     |     {req.org_name}     |     {req.prop_date}",
             size=8, colour=GREY)
    _top_border(ftr_para, CC_HEX, 4)

    # ══════════════════════════════════════════════════════════════════
    # COVER PAGE
    # ══════════════════════════════════════════════════════════════════
    _spacer(doc, 6)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run(p, "SecureITLab", size=28, bold=True, colour=NAVY)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run(p, "Cybersecurity  |  Data Privacy  |  Data Governance", size=11, colour=BLUE)
    _bottom_border(p, BLUE_HEX, 6)

    _spacer(doc, 3)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run(p, "SERVICE PROPOSAL", size=14, bold=True, colour=GREY)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run(p, req.service_name, size=18, bold=True, colour=NAVY)
    _bottom_border(p, BLUE_HEX, 6)

    _spacer(doc, 2)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run(p, "Prepared for:", size=11, colour=GREY)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run(p, req.org_name, size=16, bold=True, colour=NAVY)

    _spacer(doc, 1)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run(p, f"Proposal Date: {req.prop_date}", size=10, colour=GREY)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run(p, f"Reference: {req.prop_ref}", size=10, colour=GREY)

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════
    # S1  Organisation Overview
    # ══════════════════════════════════════════════════════════════════
    _heading(doc, "1.  Organisation Overview")
    _rule(doc)

    fields = [
        ("Organisation Name", req.org_name),
        ("Industry / Sector",  req.org_industry),
        ("Registered Address", req.org_address),
        ("Primary Contact",    req.org_contact),
        ("Email",              req.org_email),
        ("Phone",              req.org_phone),
        ("Organisation Background", req.org_bg),
        ("Engagement Objective",    req.org_obj),
    ]
    for label, val in fields:
        p = doc.add_paragraph()
        _add_run(p, f"{label}:  ", bold=True, colour=NAVY)
        _add_run(p, val or "—", colour=GREY)

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════
    # S2  About SecureITLab  (boilerplate)
    # ══════════════════════════════════════════════════════════════════
    _heading(doc, "2.  About SecureITLab")
    _rule(doc)

    _body(doc, ("SecureITLab is a remote-first cybersecurity, data privacy, and data governance "
                "consulting firm with over 20 years of industry experience. Headquartered in the "
                "Kingdom of Bahrain, with offices in the Philippines and India, SecureITLab serves "
                "clients across the Middle East, Asia-Pacific, and beyond."))
    _body(doc, ("Our team of certified professionals (CISSP, CISM, CISA, CEH, ISO 27001 LA/LI, CDPSE) "
                "delivers tailored, risk-driven solutions that empower organisations to navigate the "
                "complex landscape of cybersecurity and compliance with confidence."))
    _body(doc, "SecureITLab is certified to ISO 27001:2022, ISO 27701:2019, and ISO 22301:2019.")

    _subheading(doc, "Core Capabilities")
    caps = [
        "Consulting — Cybersecurity and Data Privacy",
        "Assurance — Penetration Testing, Vulnerability Assessment and Red Teaming",
        "Audit and Compliance — IT Audit, Cloud Audit and ISO Compliance",
        "Services Augmentation — vCISO, vCDO and Expert Resourcing",
        "Managed Services — SOC, NOC, Vulnerability Management and Take-Down",
    ]
    for c in caps:
        _bullet(doc, c)

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════
    # S3  Service Description
    # ══════════════════════════════════════════════════════════════════
    _heading(doc, "3.  Service Description")
    _rule(doc)

    p = doc.add_paragraph()
    _add_run(p, req.service_name, size=12, bold=True, colour=BLUE)
    _body(doc, req.service_desc)

    _subheading(doc, "Key Deliverables")
    for d in req.service_deliverables:
        _bullet(doc, d)

    _subheading(doc, "Methodology")
    _body(doc, ("SecureITLab employs industry-recognised methodologies, standards, and frameworks "
                "throughout this engagement. All activities are conducted in alignment with the agreed "
                "scope and subject to the client's change-control and information security policies."))
    for phase_text in proposal_data["methodology"]:
        _body(doc, phase_text)

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════
    # S4  Scope of Work
    # ══════════════════════════════════════════════════════════════════
    _heading(doc, "4.  Scope of Work")
    _rule(doc)
    _body(doc, "Define the precise boundaries of the engagement. Items not listed are considered out of scope.")

    _subheading(doc, "In Scope")
    _bullet_list(doc, req.inscope, fallback="To be defined.")

    _subheading(doc, "Out of Scope")
    _bullet_list(doc, req.outscope, fallback="To be defined.")

    _subheading(doc, "Assumptions and Dependencies")
    _bullet_list(doc, req.assumptions, fallback="To be confirmed.")

    _subheading(doc, "Client Responsibilities")
    _bullet_list(doc, req.responsibilities, fallback="To be confirmed.")

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════
    # S5  Project Timeline
    # ══════════════════════════════════════════════════════════════════
    _heading(doc, "5.  Project Timeline")
    _rule(doc)
    _body(doc, "The indicative timeline below is subject to confirmation upon contract award.")

    p = doc.add_paragraph()
    _add_run(p, "Proposed Start: ", bold=True, colour=NAVY)
    _add_run(p, req.start_date or "TBD", colour=GREY)
    _add_run(p, "     Proposed End: ", bold=True, colour=NAVY)
    _add_run(p, req.end_date or "TBD", colour=GREY)

    col_w = [Cm(3.2), Cm(6.2), Cm(2.8), Cm(5.5)]
    tbl = doc.add_table(rows=1, cols=4)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    _set_col_widths(tbl, col_w)

    hdr_cells = tbl.rows[0].cells
    for i, h in enumerate(["Phase", "Activities", "Duration", "Deliverables"]):
        _header_cell(hdr_cells[i], h, col_w[i])

    for idx, row in enumerate(timeline):
        cells = tbl.add_row().cells
        fill = LBLUE_HEX if idx % 2 == 1 else WHITE_HEX
        _data_cell(cells[0], row.phase,        col_w[0], fill)
        _data_cell(cells[1], row.activities,   col_w[1], fill)
        _data_cell(cells[2], row.duration,     col_w[2], fill)
        _data_cell(cells[3], row.deliverables, col_w[3], fill)

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════
    # S6  Investment Summary
    # ══════════════════════════════════════════════════════════════════
    _heading(doc, "6.  Investment Summary")
    _rule(doc)
    _body(doc, "All fees are exclusive of applicable taxes. Payment terms are as per the Master Service Agreement or as agreed in writing.")

    p = doc.add_paragraph()
    _add_run(p, f"Currency: {req.currency}", bold=True, colour=NAVY)

    cost_rows = req.cost_rows
    subtotal  = sum(r.total for r in cost_rows)
    tax_amt   = round(subtotal * req.tax_rate / 100, 2)
    grand     = round(subtotal + tax_amt, 2)

    ccol = [Cm(4.2), Cm(5.4), Cm(2.8), Cm(1.6), Cm(3.7)]
    ctbl = doc.add_table(rows=1, cols=5)
    ctbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    _set_col_widths(ctbl, ccol)

    ch = ctbl.rows[0].cells
    for i, h in enumerate(["Line Item", "Description", "Unit Price", "Qty", f"Total ({req.currency})"]):
        _header_cell(ch[i], h, ccol[i])

    for idx, r in enumerate(cost_rows):
        cells = ctbl.add_row().cells
        fill  = LBLUE_HEX if idx % 2 == 1 else WHITE_HEX
        _data_cell(cells[0], r.item,             ccol[0], fill)
        _data_cell(cells[1], r.desc,             ccol[1], fill)
        _data_cell(cells[2], _fmt(r.unit_price), ccol[2], fill)
        _data_cell(cells[3], str(r.qty),         ccol[3], fill)
        _data_cell(cells[4], _fmt(r.total),      ccol[4], fill)

    _summary_row(ctbl, ccol, f"Subtotal",              _fmt(subtotal), LBLUE_HEX, NAVY_HEX)
    _summary_row(ctbl, ccol, f"Tax ({req.tax_rate}%)", _fmt(tax_amt),  LBLUE_HEX, NAVY_HEX)
    _summary_row(ctbl, ccol, "TOTAL INVESTMENT",
                 f"{req.currency} {_fmt(grand)}", NAVY_HEX, WHITE_HEX, bold_label=True)

    _subheading(doc, "Payment Schedule")
    _bullet_list(doc, req.payments, fallback="As agreed.")
    _body(doc, "Validity: This proposal is valid for 30 days from the date of issue.")

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════
    # S7  Terms and Conditions  (boilerplate)
    # ══════════════════════════════════════════════════════════════════
    _heading(doc, "7.  Terms and Conditions")
    _rule(doc)
    _body(doc, ("This proposal is subject to SecureITLab's standard Master Service Agreement (MSA). "
                "Key terms are summarised below; the MSA governs in the event of any conflict."))

    tc_items = [
        ("Confidentiality",
         "Both parties agree to maintain strict confidentiality of all information exchanged. "
         "SecureITLab will not disclose client data to third parties without prior written consent."),
        ("Intellectual Property",
         "All deliverables become the property of the client upon receipt of full payment. "
         "SecureITLab retains rights to its pre-existing methodologies, tools, and frameworks."),
        ("Limitation of Liability",
         "SecureITLab's aggregate liability shall not exceed the total fees paid under this engagement. "
         "SecureITLab is not liable for indirect, consequential, or incidental damages."),
        ("Governing Law",
         "This proposal and any resulting engagement shall be governed by the laws of the Kingdom of "
         "Bahrain, unless otherwise agreed in writing."),
        ("Amendments",
         "Any change in scope, timeline, or cost requires a written Change Request signed by both "
         "parties prior to implementation."),
    ]
    for title, body in tc_items:
        p = doc.add_paragraph()
        _add_run(p, f"{title}:  ", bold=True, colour=NAVY)
        _add_run(p, body, colour=GREY)

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════
    # S8  Acceptance and Signatures
    # ══════════════════════════════════════════════════════════════════
    _heading(doc, "8.  Acceptance and Signatures")
    _rule(doc)
    _body(doc, ("By signing below, both parties confirm their agreement to the terms set out in this "
                "proposal and authorise SecureITLab to commence the engagement as described."))

    _spacer(doc, 2)

    sig_col = [Cm(7.5), Cm(2.0), Cm(7.5)]
    stbl = doc.add_table(rows=4, cols=3)
    _set_col_widths(stbl, sig_col)

    header_row = stbl.rows[0]
    _sig_header_cell(header_row.cells[0], f"Authorised Signatory — {req.org_name}", sig_col[0])
    _sig_header_cell(header_row.cells[2], "Authorised Signatory — SecureITLab",     sig_col[2])

    for row in stbl.rows[1:]:
        for cell in row.cells:
            cell.text = ""

    _spacer(doc, 2)
    _rule(doc)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run(p, "SecureITLab  |  info@secureitlab.com  |  www.secureitlab.com", size=9, colour=GREY)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run(p, "Bahrain  |  Philippines  |  India", size=9, colour=GREY)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


# ── Low-level helpers ─────────────────────────────────────────────────────────

def _add_run(para, text, size=10, bold=False, italic=False, colour=None):
    run = para.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if colour:
        run.font.color.rgb = colour
    return run

def _heading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(6)
    _add_run(p, text, size=14, bold=True, colour=NAVY)

def _subheading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(4)
    _add_run(p, text, size=11, bold=True, colour=NAVY)

def _body(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(4)
    _add_run(p, text, colour=GREY)

def _bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    _add_run(p, text, colour=GREY)

def _bullet_list(doc, items, fallback=""):
    clean = [i for i in items if i.strip()]
    if clean:
        for item in clean:
            _bullet(doc, item)
    elif fallback:
        _body(doc, fallback)

def _spacer(doc, lines=1):
    for _ in range(lines):
        doc.add_paragraph()

def _rule(doc):
    p = doc.add_paragraph()
    _bottom_border(p, BLUE_HEX, 6)

def _fmt(val):
    return f"{val:,.2f}"

def _bottom_border(para, colour_hex, size):
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), colour_hex)
    pBdr.append(bottom)
    pPr.append(pBdr)

def _top_border(para, colour_hex, size):
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    top = OxmlElement("w:top")
    top.set(qn("w:val"), "single")
    top.set(qn("w:sz"), str(size))
    top.set(qn("w:space"), "1")
    top.set(qn("w:color"), colour_hex)
    pBdr.append(top)
    pPr.append(pBdr)

def _set_col_widths(table, widths):
    tbl_el = table._tbl
    tblPr  = tbl_el.find(qn("w:tblPr")) or OxmlElement("w:tblPr")
    tblW   = OxmlElement("w:tblW")
    total  = sum(int(w.cm * 914400 / 360) for w in widths)
    tblW.set(qn("w:w"),    str(total))
    tblW.set(qn("w:type"), "dxa")
    tblPr.append(tblW)

def _twips(cm_val):
    return int(cm_val.cm * 567)

def _shade_cell(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  fill_hex)
    tcPr.append(shd)

def _cell_borders(cell, colour_hex, size):
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"),   "single")
        el.set(qn("w:sz"),    str(size))
        el.set(qn("w:color"), colour_hex)
        tcBorders.append(el)
    tcPr.append(tcBorders)

def _cell_width(cell, cm_val):
    tcPr = cell._tc.get_or_add_tcPr()
    tcW  = OxmlElement("w:tcW")
    tcW.set(qn("w:w"),    str(_twips(cm_val)))
    tcW.set(qn("w:type"), "dxa")
    tcPr.append(tcW)

def _header_cell(cell, text, width):
    cell.text = ""
    _shade_cell(cell, NAVY_HEX)
    _cell_borders(cell, BLUE_HEX, 6)
    _cell_width(cell, width)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run(p, text, size=9, bold=True, colour=WHITE)

def _data_cell(cell, text, width, fill=WHITE_HEX):
    cell.text = ""
    _shade_cell(cell, fill)
    _cell_borders(cell, CC_HEX, 4)
    _cell_width(cell, width)
    p = cell.paragraphs[0]
    _add_run(p, text or "", size=9, colour=GREY)

def _summary_row(table, col_widths, label_text, value_text, label_fill, label_colour_hex, bold_label=False):
    row   = table.add_row()
    cells = row.cells
    merged = cells[0].merge(cells[1]).merge(cells[2]).merge(cells[3])
    _shade_cell(merged, label_fill)
    _cell_borders(merged, BLUE_HEX, 6)
    p = merged.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    lc = WHITE if label_fill == NAVY_HEX else NAVY
    _add_run(p, label_text, size=10, bold=bold_label, colour=lc)

    vc = cells[4]
    _shade_cell(vc, label_fill)
    _cell_borders(vc, BLUE_HEX, 6)
    _cell_width(vc, col_widths[4])
    vp = vc.paragraphs[0]
    vc_text_colour = WHITE if label_fill == NAVY_HEX else GREY
    _add_run(vp, value_text, size=10, bold=bold_label, colour=vc_text_colour)

def _sig_header_cell(cell, text, width):
    _cell_width(cell, width)
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"),   "single")
    bottom.set(qn("w:sz"),    "6")
    bottom.set(qn("w:color"), NAVY_HEX)
    tcBorders.append(bottom)
    tcPr.append(tcBorders)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(40)
    _add_run(p, text, size=9, bold=True, colour=NAVY)