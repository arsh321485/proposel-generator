"""
update_proposals.py
Updates all 57 proposal .docx files with:
  1. Detailed phase-based methodology text (after the existing generic sentence)
  2. Filled-in Project Timeline table (Phase name, Activities, Duration, Deliverables)
"""

import sys, os, glob
sys.path.insert(0, '/sessions/great-inspiring-bardeen/mnt/SITL')

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from lxml import etree
import copy

# Import data
from proposal_data import PROPOSALS

BASE = '/sessions/great-inspiring-bardeen/mnt/SITL/Proposal_Templates'


def get_paragraph_xml_parent(para):
    return para._element.getparent()


def insert_paragraph_after(ref_para, text, bold_first_word=False):
    """Insert a new paragraph after ref_para with the given text."""
    new_para = OxmlElement('w:p')
    # Copy paragraph properties from ref_para if any
    ref_pPr = ref_para._element.find(qn('w:pPr'))
    if ref_pPr is not None:
        new_para.append(copy.deepcopy(ref_pPr))

    # Add a run with the text
    new_r = OxmlElement('w:r')
    # Copy run properties from ref_para first run if any
    ref_runs = ref_para._element.findall('.//' + qn('w:r'))
    if ref_runs:
        ref_rPr = ref_runs[0].find(qn('w:rPr'))
        if ref_rPr is not None:
            new_r.append(copy.deepcopy(ref_rPr))

    new_t = OxmlElement('w:t')
    new_t.text = text
    new_t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    new_r.append(new_t)
    new_para.append(new_r)

    # Insert after ref_para in parent
    ref_para._element.addnext(new_para)
    return new_para


def update_document(filepath, key):
    data = PROPOSALS[key]
    doc = Document(filepath)

    # --- 1. Find and update Methodology section ---
    methodology_para_idx = None
    generic_para_idx = None

    for i, para in enumerate(doc.paragraphs):
        if para.text.strip() == 'Methodology':
            methodology_para_idx = i
        if methodology_para_idx is not None and 'SecureITLab employs' in para.text:
            generic_para_idx = i
            break

    if generic_para_idx is None:
        print(f"  WARNING: Could not find methodology paragraph in {os.path.basename(filepath)}")
    else:
        # Insert phase paragraphs AFTER the generic_para, in reverse order so they end up in correct order
        ref_para = doc.paragraphs[generic_para_idx]
        for phase_text in reversed(data['methodology']):
            insert_paragraph_after(ref_para, phase_text)

    # --- 2. Update the Timeline Table (Table 0) ---
    if not doc.tables:
        print(f"  WARNING: No tables in {os.path.basename(filepath)}")
    else:
        timeline_table = doc.tables[0]
        # Verify it looks like the timeline table
        header_row = timeline_table.rows[0]
        header_texts = [c.text.strip() for c in header_row.cells]
        if 'Phase' not in header_texts[0]:
            print(f"  WARNING: Table 0 doesn't look like timeline table in {os.path.basename(filepath)}")
        else:
            phases = data['timeline']
            data_rows = timeline_table.rows[1:]  # skip header
            for i, row in enumerate(data_rows):
                if i < len(phases):
                    phase_name, activities, duration, deliverables = phases[i]
                    cells = row.cells
                    # Clear and set each cell
                    for cell, new_text in zip(cells, [phase_name, activities, duration, deliverables]):
                        # Clear existing paragraphs
                        for para in cell.paragraphs:
                            for run in para.runs:
                                run.text = ''
                        # Set text in first paragraph
                        if cell.paragraphs:
                            cell.paragraphs[0].runs[0].text = new_text if cell.paragraphs[0].runs else None
                            if not cell.paragraphs[0].runs:
                                run = cell.paragraphs[0].add_run(new_text)
                        else:
                            cell.add_paragraph(new_text)

    doc.save(filepath)
    return True


def main():
    files = sorted(glob.glob(f'{BASE}/**/*.docx', recursive=True))
    files = [f for f in files if '_test_output' not in f]

    matched = 0
    unmatched = []

    for filepath in files:
        basename = os.path.basename(filepath).replace('Proposal_', '').replace('.docx', '')
        # Find matching key in PROPOSALS
        match_key = None
        for key in PROPOSALS:
            if key == basename:
                match_key = key
                break

        if match_key is None:
            unmatched.append(basename)
            continue

        try:
            update_document(filepath, match_key)
            print(f"  OK: {basename}")
            matched += 1
        except Exception as e:
            print(f"  ERROR: {basename}: {e}")

    print(f"\nDone: {matched} updated, {len(unmatched)} unmatched")
    if unmatched:
        print("Unmatched:", unmatched)


if __name__ == '__main__':
    main()
